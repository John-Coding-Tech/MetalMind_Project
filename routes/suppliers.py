"""
routes/suppliers.py — Save, list, and AI-search saved suppliers.
"""

import hashlib
import json
import logging
import mimetypes
import re
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from typing import Literal
from urllib.parse import quote as url_quote

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from db import get_db
from models import SavedSupplier, SupplierAttachment
from engine.ai_engine import call_model
from services.tavily_client import enrich_url, is_available as tavily_available
from services.serper_client import search as serper_search, SerperError
from config import MAX_UPLOAD_BYTES, ALLOWED_UPLOAD_EXTS


# Filesystem location for uploaded files. Resolved relative to project root
# so it works identically in dev and on Railway. Created lazily on first upload.
_PROJECT_ROOT = Path(__file__).resolve().parent.parent
_UPLOAD_ROOT  = _PROJECT_ROOT / "uploads"

# Filename sanitiser — keep original extension, strip path traversal and
# control characters. Uniqueness is handled by prefixing with attachment id.
_SAFE_FILENAME = re.compile(r"[^\w.\-]+", re.UNICODE)


def _sanitize_filename(name: str) -> str:
    name = Path(name).name   # strip any path components
    name = _SAFE_FILENAME.sub("_", name).strip("._")
    return name or "file"

_log = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["suppliers"])


def _get_active_supplier(db: Session, supplier_id: int) -> SavedSupplier:
    """Return a non-deleted supplier or raise 404."""
    supplier = (db.query(SavedSupplier)
                  .filter(SavedSupplier.id == supplier_id, SavedSupplier.is_saved == True)  # noqa: E712
                  .first())
    if not supplier:
        raise HTTPException(status_code=404, detail="Supplier not found")
    return supplier


class SaveSupplierRequest(BaseModel):
    supplier_name: str = Field(max_length=500)
    country: str | None = Field(None, max_length=100)
    price_display: str | None = Field(None, max_length=100)
    price_usd: float | None = Field(None, ge=0)
    risk_level: str | None = Field(None, max_length=50)
    risk_score: float | None = Field(None, ge=0, le=1)
    risk_reasons: list[str] | None = None
    value_score: float | None = Field(None, ge=0)
    url: str | None = Field(None, max_length=1000)
    description: str | None = Field(None, max_length=5000)
    trust: str | None = Field(None, max_length=50)
    anomalies: dict | None = None
    ai_adjustment: dict | None = None


class ChatTurn(BaseModel):
    role: Literal["user", "assistant"]
    content: str = Field(max_length=10000)


class AiSearchRequest(BaseModel):
    query: str = Field(max_length=2000)
    selected_ids: list[int] | None = Field(None, max_length=100)
    history: list[ChatTurn] | None = Field(None, max_length=20)


class UpdateNoteRequest(BaseModel):
    notes: str = Field(max_length=50000)


class AssessmentUpdateRequest(BaseModel):
    # Tier 1
    decision_stage: str | None = Field(None, max_length=50)
    rating: int | None = Field(None, ge=0, le=5)
    tags: list[str] | None = None
    pros: list[str] | None = None
    cons: list[str] | None = None
    # Tier 2
    quoted_price: float | None = Field(None, ge=0)
    quoted_currency: str | None = Field(None, max_length=3)
    quoted_unit: str | None = Field(None, max_length=20)
    moq: int | None = Field(None, ge=0)
    lead_time_days: int | None = Field(None, ge=0, le=3650)
    payment_terms: str | None = Field(None, max_length=500)
    incoterms: str | None = Field(None, max_length=10)
    # Tier 3
    sample_status: str | None = Field(None, max_length=30)
    sample_quality: int | None = Field(None, ge=0, le=5)
    factory_verified_via: list[str] | None = None
    # Legacy verification columns (kept for backward compat; prefer reference_1/2/3 for new data)
    coating_confirmed: str | None = Field(None, max_length=20)
    core_material_confirmed: str | None = Field(None, max_length=20)
    fire_rating_confirmed: str | None = Field(None, max_length=10)
    # Free-text reference fields (shown as Reference 1/2/3 in the UI)
    reference_1: str | None = Field(None, max_length=2000)
    reference_2: str | None = Field(None, max_length=2000)
    reference_3: str | None = Field(None, max_length=2000)
    warranty_years: int | None = Field(None, ge=0, le=100)
    next_action_date: str | None = None  # ISO "YYYY-MM-DD"
    # Free-form notes — shared with existing endpoint
    notes: str | None = Field(None, max_length=50000)


@router.post("/save-supplier")
def save_supplier(req: SaveSupplierRequest, db: Session = Depends(get_db)):
    existing = db.query(SavedSupplier).filter(
        SavedSupplier.supplier_name == req.supplier_name,
        SavedSupplier.url == req.url,
    ).first()

    if existing:
        if existing.is_saved:
            raise HTTPException(status_code=409, detail="Supplier already saved")
        # Soft-deleted row: reactivate it, refresh the system-scored fields
        # with the latest run, but KEEP user-entered Details (rating, pros,
        # cons, quoted_price, moq, payment_terms, sample_status, etc.)
        existing.is_saved     = True
        existing.country      = req.country
        existing.price_display = req.price_display
        existing.price_usd    = req.price_usd
        existing.risk_level   = req.risk_level
        existing.risk_score   = req.risk_score
        existing.risk_reasons = req.risk_reasons
        existing.value_score  = req.value_score
        existing.description  = req.description
        existing.trust        = req.trust
        existing.anomalies    = req.anomalies
        existing.ai_adjustment = req.ai_adjustment
        db.commit()
        db.refresh(existing)
        return _to_dict(existing)

    supplier = SavedSupplier(
        supplier_name=req.supplier_name,
        country=req.country,
        price_display=req.price_display,
        price_usd=req.price_usd,
        risk_level=req.risk_level,
        risk_score=req.risk_score,
        risk_reasons=req.risk_reasons,
        value_score=req.value_score,
        url=req.url,
        description=req.description,
        trust=req.trust,
        anomalies=req.anomalies,
        ai_adjustment=req.ai_adjustment,
    )
    db.add(supplier)
    db.commit()
    db.refresh(supplier)
    return _to_dict(supplier)


@router.get("/saved-suppliers")
def list_saved(db: Session = Depends(get_db)):
    # Only return currently-saved rows; soft-deleted rows stay in the DB
    # so their user-filled Details survive re-save.
    suppliers = (db.query(SavedSupplier)
                   .filter(SavedSupplier.is_saved == True)   # noqa: E712
                   .order_by(SavedSupplier.saved_at.desc())
                   .all())
    return [_to_dict(s) for s in suppliers]


@router.delete("/saved-supplier/{supplier_id}")
def delete_saved(supplier_id: int, db: Session = Depends(get_db)):
    """Soft-delete: flip is_saved to False. The row and its attachments
    stay on disk/DB so a subsequent re-save restores the user's Details."""
    supplier = _get_active_supplier(db, supplier_id)
    supplier.is_saved = False
    db.commit()
    return {"ok": True}


@router.post("/supplier-report/{supplier_id}")
def supplier_report(supplier_id: int, refresh: bool = False, db: Session = Depends(get_db)):
    """Generate a deep-dive report for one supplier using web search + AI.
    Uses cached report unless refresh=true."""
    s = _get_active_supplier(db, supplier_id)

    # Return cached report if available and not forcing refresh
    if not refresh and s.deep_report:
        cached = dict(s.deep_report)
        cached["cached"] = True
        cached["generated_at"] = s.report_generated_at.isoformat() if s.report_generated_at else None
        return cached

    # Gather web search results on multiple dimensions
    searches = {
        "overview":      f"{s.supplier_name} ACP aluminium composite panel company history background",
        "capacity":      f"{s.supplier_name} production capacity factory size employees",
        "certifications": f"{s.supplier_name} ISO certification quality standards",
        "reputation":    f"{s.supplier_name} customer reviews reputation feedback",
        "news":          f"{s.supplier_name} news 2025 2026",
    }
    search_context = {}
    for topic, q in searches.items():
        try:
            hits = serper_search(q, max_results=5)
            if hits:
                search_context[topic] = "\n".join([
                    f"- {h.get('title', '')}: {h.get('content', '')[:300]}"
                    for h in hits[:5]
                ])
        except Exception as e:
            _log.warning("Report search [%s] failed for %s: %s", topic, s.supplier_name, e)

    # Also enrich the supplier's website for direct facts
    website_content = ""
    if s.url and tavily_available():
        try:
            result = enrich_url(s.url)
            if result and result.get("content"):
                website_content = result["content"][:6000]
        except Exception as e:
            _log.warning("Website enrichment failed for %s: %s", s.supplier_name, e)

    if not search_context and not website_content:
        return {
            "supplier_name": s.supplier_name,
            "sections": {},
            "summary": "Unable to gather enough information for a deep-dive report. Web search and website fetching both failed.",
            "confidence": 0.0,
        }

    # Build the prompt
    search_block = "\n\n".join([
        f"=== {topic.upper()} SEARCH ===\n{content}"
        for topic, content in search_context.items()
    ])
    website_block = f"\n\n=== WEBSITE CONTENT ({s.url}) ===\n{website_content}" if website_content else ""

    saved_profile = (
        f"Name: {s.supplier_name}\n"
        f"Country: {s.country}\n"
        f"Price: {s.price_display}\n"
        f"Risk: {s.risk_level} (score: {s.risk_score})\n"
        f"Value Score: {s.value_score}/100\n"
        f"Trust: {s.trust}\n"
        f"URL: {s.url}\n"
        f"Risk reasons: {', '.join(s.risk_reasons) if s.risk_reasons else 'none'}\n"
        f"Notes: {s.notes or 'none'}"
    )

    prompt = f"""You are a supplier due diligence analyst. Generate a comprehensive deep-dive report.

SAVED PROFILE:
{saved_profile}

{search_block}{website_block}

Produce a JSON object ONLY (no markdown) with this exact shape:

{{
  "supplier_name": "{s.supplier_name}",
  "sections": {{
    "overview": "2-3 sentences about company history, founding, and main products. Cite specific facts from sources.",
    "capacity": "Production capacity, factory size, employee count if known. Say 'not found' if unavailable.",
    "certifications": "Specific certifications (ISO, etc.) with years if known.",
    "reputation": "Customer reviews and reputation signals. Be specific about sources.",
    "recent_news": "Any recent developments, contracts, expansions, issues (2024-2026).",
    "risk_signals": "Red flags, concerns, or notable risks for a buyer."
  }},
  "summary": "Executive summary: 3-5 sentences giving an actionable verdict for a buyer considering this supplier.",
  "confidence": 0.0-1.0
}}

Rules:
- Be specific — cite facts from the search results / website content. Never invent.
- If a section has no info, write "Not found in available sources."
- Reply in English unless the supplier name suggests another locale.
- Output ONLY the JSON object. No markdown, no prose."""

    try:
        raw = call_model(prompt)
        if not raw:
            return {"supplier_name": s.supplier_name, "error": "AI model returned empty response"}
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        report = json.loads(cleaned)
        # Cache in DB
        from datetime import datetime, timezone
        s.deep_report = report
        s.report_generated_at = datetime.now(timezone.utc)
        db.commit()
        report["cached"] = False
        report["generated_at"] = s.report_generated_at.isoformat()
        return report
    except json.JSONDecodeError:
        _log.warning("Report AI response unparseable for %s", s.supplier_name)
        return {"supplier_name": s.supplier_name, "error": "Could not parse AI response"}
    except Exception:
        _log.exception("Report generation failed for %s", s.supplier_name)
        return {"supplier_name": s.supplier_name, "error": "Report generation failed"}


@router.patch("/saved-supplier/{supplier_id}/notes")
def update_notes(supplier_id: int, req: UpdateNoteRequest, db: Session = Depends(get_db)):
    supplier = _get_active_supplier(db, supplier_id)
    supplier.notes = req.notes
    db.commit()
    db.refresh(supplier)
    return _to_dict(supplier)


@router.patch("/saved-supplier/{supplier_id}/assessment")
def update_assessment(supplier_id: int, req: AssessmentUpdateRequest, db: Session = Depends(get_db)):
    """Update any subset of the user's primary-decision-data fields.
    Auto-save friendly: only fields present in the request body are touched."""
    supplier = _get_active_supplier(db, supplier_id)
    data = req.model_dump(exclude_unset=True)
    if "next_action_date" in data and data["next_action_date"]:
        from datetime import date as _date
        try:
            data["next_action_date"] = _date.fromisoformat(data["next_action_date"])
        except ValueError:
            raise HTTPException(status_code=400, detail="next_action_date must be YYYY-MM-DD")
    for k, v in data.items():
        setattr(supplier, k, v)
    db.commit()
    db.refresh(supplier)
    return _to_dict(supplier)


# =========================================================================
# Ask-AI — Step 2: data feeding + smart trigger
#
# The prompt exposes each supplier in up to four data layers, ordered by
# authority. User-entered data is labelled VERIFIED (not "primary") because
# LLMs weight that word as ground truth more reliably.
#
#     [VERIFIED]     user-filled Tier 1/2/3 fields         — absolute authority
#     [SECONDARY-A]  rule-based scores + anomalies         — deterministic
#     [SECONDARY-B]  AI-derived adjustments (optional)     — lowest authority
#     [LIVE-WEB]     Tavily/Serper fetched this turn       — fresh but unverified
#
# Layer-C (live web) is triggered ONLY when the VERIFIED data can't answer
# the query — avoiding wasted API calls while guaranteeing accuracy.
# =========================================================================

# Query keyword → VERIFIED field on SavedSupplier.
# Sentinel values:
#   None        — DB does not cover this dimension (always search the web).
#   "__enrich__" — trigger Tavily URL extract (contact info from the site).
_QUERY_DIMENSIONS: list[tuple[list[str], str | None]] = [
    # Commercial facts (user-filled in Edit Details)
    (["moq", "minimum order", "最小起订"],                             "moq"),
    (["lead time", "delivery time", "交期", "发货", "ship", "dispatch"],"lead_time_days"),
    (["payment", "terms", "付款", "t/t", "lc", "letter of credit"],    "payment_terms"),
    (["incoterm", "fob", "cif", "exw", "ddp"],                         "incoterms"),
    (["warranty", "保修", "guarantee"],                                 "warranty_years"),
    (["quoted price", "quote", "cost", "报价"],                         "quoted_price"),
    (["currency"],                                                      "quoted_currency"),
    # Verification / sample
    (["sample"],                                                        "sample_status"),
    (["factory visit", "verified", "audit", "验厂"],                     "factory_verified_via"),
    # User's own assessment
    (["rating", "star", "评分"],                                        "rating"),
    (["decision", "stage", "shortlist"],                                "decision_stage"),
    (["pros", "strength", "优点"],                                      "pros"),
    (["cons", "weakness", "downside", "缺点"],                          "cons"),
    (["tag"],                                                           "tags"),
    (["next action", "follow up", "deadline"],                          "next_action_date"),
    (["note", "comment", "备注"],                                       "notes"),
    # Contact info → enrich the supplier's website
    (["contact", "email", "phone", "address", "reach", "call",
      "location", "where", "whatsapp", "wechat",
      "联系", "电话", "邮箱", "地址", "微信"],                           "__enrich__"),
    # DB does not cover these — always search
    (["capacity", "产能", "production", "output"],                      None),
    (["review", "reputation", "feedback", "评价", "口碑"],              None),
    (["founded", "established", "history", "year", "成立"],             None),
    (["certification", "certified", "iso", "认证", "证书"],             None),
    (["news", "recent", "latest", "新闻"],                              None),
    (["workforce", "employee", "staff", "员工"],                        None),
]

# Override keyword lists — force web search even when VERIFIED is present.
_VERIFY_KEYWORDS    = ["verify", "confirm", "double-check", "is it true",
                       "really", "actually", "check", "核实", "确认"]
_FRESHNESS_KEYWORDS = ["latest", "current", "now", "today", "still", "recent",
                       "最新", "现在", "目前"]

# Which anomaly/adjustment reason words cast doubt on which VERIFIED field.
_DIMENSION_CONCERN_WORDS: dict[str, list[str]] = {
    "quoted_price":   ["price", "cheap", "low", "expensive", "suspicious", "outlier"],
    "moq":            ["moq", "order"],
    "lead_time_days": ["lead", "delivery"],
    "warranty_years": ["warranty"],
}


def _match_dimensions(query: str) -> list[tuple[list[str], str | None]]:
    """Return the subset of _QUERY_DIMENSIONS whose keywords appear in query."""
    ql = query.lower()
    return [(kw, field) for kw, field in _QUERY_DIMENSIONS if any(k in ql for k in kw)]


def _is_placeholder(val) -> bool:
    """True when a stored VERIFIED value is effectively empty."""
    if val is None:
        return True
    if isinstance(val, (list, tuple, dict)):
        return len(val) == 0
    if isinstance(val, str):
        return val.strip().lower() in ("", "unknown", "n/a", "na", "tbd", "?", "none", "-")
    if isinstance(val, (int, float)):
        return val == 0
    return False


def _dimension_is_flagged(s: SavedSupplier, field_name: str) -> bool:
    """True if anomalies or ai_adjustment cast doubt on this dimension."""
    words = _DIMENSION_CONCERN_WORDS.get(field_name)
    if not words:
        return False
    for a in ((s.anomalies or {}).get("anomalies") or []):
        if any(w in a.lower() for w in words):
            return True
    reason = ((s.ai_adjustment or {}).get("reason") or "").lower()
    if reason and any(w in reason for w in words):
        return True
    return False


def _primary_is_reliable(s: SavedSupplier, field_name: str, query_lower: str) -> bool:
    """
    Return True when the VERIFIED value for `field_name` on `s` is trustworthy
    enough to answer the query without a web lookup.

    Returns False (→ search this supplier) when ANY of:
      1. Value is a placeholder/empty.
      2. Query asks to verify/confirm.
      3. Query needs freshness (latest / current / now).
      4. anomalies or ai_adjustment already flagged this dimension.
    """
    if _is_placeholder(getattr(s, field_name, None)):
        return False
    if any(kw in query_lower for kw in _VERIFY_KEYWORDS):
        return False
    if any(kw in query_lower for kw in _FRESHNESS_KEYWORDS):
        return False
    if _dimension_is_flagged(s, field_name):
        return False
    return True


# ------------------------------------------------------------------
# Four-layer supplier block formatters
# ------------------------------------------------------------------

def _fmt_verified(s: SavedSupplier) -> str:
    """[VERIFIED] block — user-filled Tier 1/2/3 fields. Skips empty values."""
    lines = ["[VERIFIED — entered by user after direct verification (quote / sample / factory check)]"]

    if s.decision_stage: lines.append(f"  Decision stage: {s.decision_stage}")
    if s.rating:         lines.append(f"  Rating: {s.rating}/5")
    if s.tags:           lines.append(f"  Tags: {', '.join(s.tags)}")
    if s.pros:           lines.append(f"  Pros: {'; '.join(s.pros)}")
    if s.cons:           lines.append(f"  Cons: {'; '.join(s.cons)}")

    if s.quoted_price is not None:
        head = f"  Quoted price: {s.quoted_price} {s.quoted_currency or 'USD'}/{s.quoted_unit or 'sqm'}"
        extras = []
        if s.moq is not None:            extras.append(f"MOQ: {s.moq}")
        if s.lead_time_days is not None: extras.append(f"Lead time: {s.lead_time_days} days")
        if extras: head += "  (" + ", ".join(extras) + ")"
        lines.append(head)
    else:
        if s.moq is not None:            lines.append(f"  MOQ: {s.moq}")
        if s.lead_time_days is not None: lines.append(f"  Lead time: {s.lead_time_days} days")
    if s.payment_terms: lines.append(f"  Payment terms: {s.payment_terms}")
    if s.incoterms:     lines.append(f"  Incoterms: {s.incoterms}")

    if s.sample_status:
        sq = f"  Sample status: {s.sample_status}"
        if s.sample_quality is not None: sq += f" (quality: {s.sample_quality}/5)"
        lines.append(sq)
    if s.factory_verified_via:
        lines.append(f"  Factory verified via: {', '.join(s.factory_verified_via)}")
    if s.warranty_years is not None: lines.append(f"  Warranty: {s.warranty_years} years")

    # References (three free-text fields shown as "Reference 1/2/3" in the UI)
    refs = [r for r in (s.reference_1, s.reference_2, s.reference_3) if r]
    if refs: lines.append(f"  References: {'; '.join(refs)}")

    if s.next_action_date: lines.append(f"  Next action: {s.next_action_date}")
    if s.notes:            lines.append(f"  Notes: {s.notes}")

    if len(lines) == 1:
        lines.append("  (no user data filled yet)")
    return "\n".join(lines)


def _fmt_secondary_a(s: SavedSupplier) -> str:
    """Rule-based block: risk, value, description, anomalies."""
    lines = ["[SECONDARY-A — rule-based, deterministic snapshot]"]
    if s.risk_level:
        rs = f"  Risk: {s.risk_level}"
        if s.risk_score is not None:
            rs += f" (score: {s.risk_score:.2f})"
        lines.append(rs)
    if s.risk_reasons:
        lines.append(f"  Risk reasons: {'; '.join(s.risk_reasons)}")
    if s.value_score is not None:
        lines.append(f"  Value score: {s.value_score}/100")
    if s.price_display:
        lines.append(f"  Displayed price: {s.price_display}")
    if s.description:
        lines.append(f"  Description: {s.description[:300]}")

    anoms    = (s.anomalies or {}).get("anomalies") or []
    severity = (s.anomalies or {}).get("severity", "none")
    if anoms:
        lines.append(f"  Anomalies [severity: {severity}]:")
        for a in anoms:
            lines.append(f"    - {a}")
    return "\n".join(lines)


def _fmt_secondary_b(s: SavedSupplier) -> str:
    """AI-derived adjustments. Returns '' when nothing meaningful to show."""
    adj = s.ai_adjustment or {}
    if not adj or adj.get("confidence", 0) < 0.5 or adj.get("adjustment", 0) == 0:
        return ""
    sign = "+" if adj["adjustment"] > 0 else ""
    line = (f"  AI adjustment: {sign}{adj['adjustment']} "
            f"(confidence {adj['confidence']:.2f}, "
            f"reason: \"{adj.get('reason', '')}\")")
    return "[SECONDARY-B — AI-derived, lower authority]\n" + line


# ── Deep-Research integration (Strategy B: facts only, no AI opinions) ──
# Deep Report is AI-synthesized from web data, so echoing its opinion sections
# back into Ask AI would create an "AI answering AI" loop. We only pipe the
# fact-heavy sections (certifications, company overview, capacity). Opinion
# sections (summary / reputation / risk_signals) are DELIBERATELY SKIPPED.

_DEEP_REPORT_SAFE_SECTIONS = ["certifications", "company_overview", "capacity"]

# Maps _QUERY_DIMENSIONS keywords (the ones with field=None) to the deep_report
# section that can satisfy them, so we can skip live web searches when the
# supplier's cached report already covers it.
_DEEP_REPORT_SECTION_MAP: dict[str, str] = {
    "certification": "certifications",
    "certified":     "certifications",
    "iso":           "certifications",
    "认证":          "certifications",
    "证书":          "certifications",
    "capacity":      "capacity",
    "production":    "capacity",
    "output":        "capacity",
    "产能":          "capacity",
    "founded":       "company_overview",
    "established":   "company_overview",
    "history":       "company_overview",
    "year":          "company_overview",
    "成立":          "company_overview",
}

# How fresh the cached report needs to be before we trust it enough to
# suppress a live web search. Beyond this, we include the report in the
# prompt as context but still search fresh web data.
_DEEP_REPORT_FRESH_DAYS = 30


def _deep_report_has_content(s: SavedSupplier, section: str) -> bool:
    """True when the supplier has a non-empty, non-'not found' entry for
    the given deep_report section."""
    dr = s.deep_report or {}
    content = (dr.get("sections", {}).get(section) or "").strip()
    if not content:
        return False
    low = content.lower()
    return not ("not found" in low or "not available" in low)


def _deep_report_is_fresh(s: SavedSupplier) -> bool:
    """True when the report was generated within _DEEP_REPORT_FRESH_DAYS."""
    ts = s.report_generated_at
    if ts is None:
        return False
    try:
        from datetime import datetime, timezone, timedelta
        now = datetime.now(timezone.utc)
        # stored_at may be naive — normalize to UTC for a safe subtract
        if ts.tzinfo is None:
            ts = ts.replace(tzinfo=timezone.utc)
        return (now - ts) < timedelta(days=_DEEP_REPORT_FRESH_DAYS)
    except Exception:
        return False


def _deep_report_section_for_keywords(keywords: list[str]) -> str | None:
    """Return the safe section name that answers these query keywords, or None."""
    for kw in keywords:
        section = _DEEP_REPORT_SECTION_MAP.get(kw)
        if section and section in _DEEP_REPORT_SAFE_SECTIONS:
            return section
    return None


def _fmt_deep_research(s: SavedSupplier) -> str:
    """
    [DEEP-RESEARCH] block — FACTS ONLY (Strategy B).
    Returns "" when the supplier has no report or no usable fact sections.
    """
    dr = s.deep_report
    if not dr:
        return ""

    sections = dr.get("sections", {})
    gen_at = s.report_generated_at.isoformat()[:10] if s.report_generated_at else "?"

    fact_lines: list[str] = []
    for key in _DEEP_REPORT_SAFE_SECTIONS:
        val = (sections.get(key) or "").strip()
        if not val:
            continue
        if "not found" in val.lower() or "not available" in val.lower():
            continue
        label = key.replace("_", " ").title()
        fact_lines.append(f"  {label}: {val[:400]}")

    if not fact_lines:
        return ""

    header = (
        f"[DEEP-RESEARCH — AI-synthesized from web data (generated {gen_at}).\n"
        f"                 Only concrete facts shown below; AI-opinion\n"
        f"                 sections (summary, reputation, risk_signals)\n"
        f"                 deliberately omitted to avoid AI-echo.]"
    )
    return header + "\n" + "\n".join(fact_lines)


def _fmt_live_web(name: str, enriched_data: dict, web_results: dict) -> str:
    """Live-web block. Only rendered when something was actually fetched."""
    chunks = []
    if name in enriched_data:
        chunks.append(f"  Website content [source: web-fetch]:\n    {enriched_data[name][:2000]}")
    if name in web_results:
        chunks.append(f"  Google search results [source: web-search]:\n    {web_results[name][:1500]}")
    if not chunks:
        return ""
    return "[LIVE-WEB — fetched this turn, may be stale after a few minutes]\n" + "\n\n".join(chunks)


def _format_supplier_block(
    s: SavedSupplier,
    enriched_data: dict,
    web_results: dict,
    attachments: list[SupplierAttachment] | None = None,
) -> str:
    """Assemble a single supplier's four-layer block (plus attachments if any)."""
    header = (f"==============================\n"
              f"SUPPLIER: {s.supplier_name} ({s.country or 'unknown'})\n"
              f"URL: {s.url or '(none)'}\n"
              f"==============================")
    parts = [header, _fmt_verified(s), _fmt_secondary_a(s)]
    sb = _fmt_secondary_b(s)
    if sb:
        parts.append(sb)
    at = _fmt_attachments(s, attachments or [])
    if at:
        parts.append(at)
    # DEEP-RESEARCH block comes BEFORE LIVE-WEB in the prompt because
    # fresh LIVE-WEB (when present) should take priority, and the RULES
    # tell the AI so. The block itself only lists concrete facts.
    dr = _fmt_deep_research(s)
    if dr:
        parts.append(dr)
    lw = _fmt_live_web(s.supplier_name, enriched_data, web_results)
    if lw:
        parts.append(lw)
    return "\n\n".join(parts)


# =========================================================================
# Layer-C fetch infrastructure: parallel execution + TTL cache
# =========================================================================

# Separate caches per fetch type. Keys are normalized so the same underlying
# lookup from different phrasings hits the same entry.
_LAYER_C_TTL = 300   # 5 minutes
_LAYER_C_MAX_WORKERS = 8
_enrich_cache: dict[str, tuple[float, str | None]] = {}   # url → (ts, content)
_search_cache: dict[str, tuple[float, str | None]] = {}   # key → (ts, formatted block)


def _cache_get(cache: dict, key: str):
    hit = cache.get(key)
    if hit is None:
        return None
    ts, data = hit
    if time.time() - ts > _LAYER_C_TTL:
        cache.pop(key, None)
        return None
    return data


def _cache_set(cache: dict, key: str, data) -> None:
    cache[key] = (time.time(), data)


def _serper_cache_key(supplier_name: str, query: str) -> str:
    # Collapse whitespace + lowercase so "iso cert" and "ISO  cert" share a key.
    norm = " ".join(query.lower().split())
    digest = hashlib.sha1(f"{supplier_name}||{norm}".encode("utf-8")).hexdigest()
    return digest[:16]


def _fetch_enrich_one(s: SavedSupplier) -> tuple[str, str | None]:
    """Run (or return cached) Tavily extract for one supplier."""
    if not s.url:
        return (s.supplier_name, None)

    cached = _cache_get(_enrich_cache, s.url)
    if cached is not None:
        _log.info("[ai-search] enrich CACHE HIT %s", s.url[:60])
        return (s.supplier_name, cached)

    try:
        result = enrich_url(s.url)
        content = (result or {}).get("content") or None
        if content:
            content = content[:4000]
        _cache_set(_enrich_cache, s.url, content)
        return (s.supplier_name, content)
    except Exception as e:
        _log.warning("Enrichment failed for %s: %s", s.supplier_name, e)
        return (s.supplier_name, None)


def _fetch_serper_one(s: SavedSupplier, user_query: str) -> tuple[str, str | None]:
    """Run (or return cached) Serper search for one supplier + user query."""
    cache_key = _serper_cache_key(s.supplier_name, user_query)

    cached = _cache_get(_search_cache, cache_key)
    if cached is not None:
        _log.info("[ai-search] serper CACHE HIT %s", s.supplier_name[:40])
        return (s.supplier_name, cached)

    try:
        q = f"{s.supplier_name} {user_query}"
        hits = serper_search(q, max_results=5)
        if not hits:
            _cache_set(_search_cache, cache_key, None)
            return (s.supplier_name, None)
        block = "\n".join(
            f"- {h.get('title', '')}: {h.get('content', '')[:300]}"
            for h in hits[:5]
        )
        _cache_set(_search_cache, cache_key, block)
        return (s.supplier_name, block)
    except (SerperError, Exception) as e:
        _log.warning("Web search failed for %s: %s", s.supplier_name, e)
        return (s.supplier_name, None)


def _fetch_layer_c_parallel(
    enrich_targets: dict[int, SavedSupplier],
    search_targets: dict[int, SavedSupplier],
    user_query: str,
) -> tuple[dict[str, str], dict[str, str]]:
    """
    Run all Tavily enrichments and Serper searches concurrently.
    Wall-clock time drops from sum-of-calls to roughly the slowest call.
    """
    enriched_data: dict[str, str] = {}
    web_results:   dict[str, str] = {}

    jobs: list = []
    tavily_on = tavily_available()

    with ThreadPoolExecutor(max_workers=_LAYER_C_MAX_WORKERS) as ex:
        if tavily_on:
            for s in enrich_targets.values():
                jobs.append(("enrich", ex.submit(_fetch_enrich_one, s)))
        for s in search_targets.values():
            jobs.append(("search", ex.submit(_fetch_serper_one, s, user_query)))

        for kind, fut in jobs:
            name, content = fut.result()
            if not content:
                continue
            if kind == "enrich":
                enriched_data[name] = content
            else:
                web_results[name] = content

    return enriched_data, web_results


# =========================================================================
# Attachments → AI readable text
#
# Attachments are stored on disk by `POST /suppliers/{id}/attachments`.
# For Ask-AI we extract readable text from text-based formats only
# (PDF, Word, Excel, CSV, plain text). Images / archives are listed by
# name so the AI knows they exist but cannot read them.
# =========================================================================

_ATTACHMENT_TEXT_BUDGET_PER_FILE = 3000    # chars per file to keep prompt sane
_ATTACHMENT_TEXT_BUDGET_TOTAL    = 8000    # chars per supplier total

# Cache extracted text per (path, mtime) so we don't re-parse unchanged files.
_attachment_text_cache: dict[tuple[str, float], str] = {}


def _extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF; swallows errors and returns '' on failure."""
    try:
        import fitz  # pymupdf
    except ImportError:
        return ""
    try:
        out = []
        with fitz.open(str(path)) as doc:
            for page in doc:
                out.append(page.get_text("text"))
                if sum(len(x) for x in out) > _ATTACHMENT_TEXT_BUDGET_PER_FILE * 2:
                    break
        return "\n".join(out).strip()
    except Exception as e:
        _log.warning("PDF extract failed for %s: %s", path.name, e)
        return ""


def _extract_docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError:
        return ""
    try:
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        # Include table cells — many supplier quotes live in Word tables
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        return "\n".join(paras).strip()
    except Exception as e:
        _log.warning("DOCX extract failed for %s: %s", path.name, e)
        return ""


def _extract_xlsx_text(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return ""
    try:
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        out = []
        for sheet in wb.worksheets:
            out.append(f"== Sheet: {sheet.title} ==")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    out.append(" | ".join(cells))
                if sum(len(x) for x in out) > _ATTACHMENT_TEXT_BUDGET_PER_FILE * 2:
                    break
        wb.close()
        return "\n".join(out).strip()
    except Exception as e:
        _log.warning("XLSX extract failed for %s: %s", path.name, e)
        return ""


def _extract_plain_text(path: Path) -> str:
    try:
        with path.open("r", encoding="utf-8", errors="replace") as f:
            return f.read().strip()
    except Exception as e:
        _log.warning("Plain text read failed for %s: %s", path.name, e)
        return ""


# ── Image attachments → Gemma vision ────────────────────────────────────
# Images aren't turned into text here; we hand the raw pixels to Gemma's
# multimodal endpoint. Resize + JPEG-compress first so a 4K screenshot
# doesn't blow up the request size.

_MAX_IMAGE_DIMENSION   = 1600  # px — bigger gets downscaled
_JPEG_QUALITY          = 85
_MAX_IMAGES_PER_SUPPLIER = 3   # cap vision tokens; newest first

# Cache prepared JPEG bytes per (path, mtime) so repeat queries within the
# Layer-C TTL window don't re-read + re-compress the same image.
_image_bytes_cache: dict[tuple[str, float], bytes | None] = {}


def _is_image_attachment(att: SupplierAttachment) -> bool:
    name = (att.filename or "").lower()
    mime = (att.mime_type or "").lower()
    return mime.startswith("image/") or name.endswith(
        (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp")
    )


def _prepare_image_for_gemma(att: SupplierAttachment) -> bytes | None:
    """
    Load, resize (long edge ≤ 1600px), and JPEG-compress an image attachment.
    Returns JPEG bytes, or None on any failure. Cached per (path, mtime).
    """
    disk_path = _PROJECT_ROOT / att.stored_path
    if not disk_path.exists():
        return None
    try:
        mtime = disk_path.stat().st_mtime
    except OSError:
        return None

    cache_key = (str(disk_path), mtime)
    if cache_key in _image_bytes_cache:
        return _image_bytes_cache[cache_key]

    try:
        from PIL import Image
        import io
        with Image.open(disk_path) as img:
            # Strip alpha for JPEG output; paste on white to avoid black halos
            if img.mode in ("RGBA", "LA", "P"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                bg.paste(img.convert("RGBA"), mask=img.convert("RGBA").split()[-1])
                img = bg
            elif img.mode != "RGB":
                img = img.convert("RGB")

            w, h = img.size
            longest = max(w, h)
            if longest > _MAX_IMAGE_DIMENSION:
                scale = _MAX_IMAGE_DIMENSION / longest
                img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=_JPEG_QUALITY, optimize=True)
            data = buf.getvalue()
    except Exception as e:
        _log.warning("Image prep failed for %s: %s", att.filename, e)
        _image_bytes_cache[cache_key] = None
        return None

    _image_bytes_cache[cache_key] = data
    return data


def _collect_images_for_ai(
    atts_by_supplier: dict[int, list[SupplierAttachment]],
    max_images: int = _MAX_IMAGES_PER_SUPPLIER,
) -> tuple[list[bytes], list[tuple[str, str]]]:
    """
    Walk the attachment map and prepare up to `max_images` JPEG payloads per
    supplier. Returns:
        images:    list[bytes]  — in the order the AI will see them
        vision_index: list[(supplier_name, filename)] — parallel to images,
                                                        so the prompt can
                                                        label each thumbnail
    """
    images: list[bytes] = []
    vision_index: list[tuple[str, str]] = []

    for supplier_id, atts in atts_by_supplier.items():
        image_atts = [a for a in atts if _is_image_attachment(a)]
        if not image_atts:
            continue
        # atts are already ordered newest-first by the DB query
        for a in image_atts[:max_images]:
            data = _prepare_image_for_gemma(a)
            if not data:
                continue
            images.append(data)
            # supplier_name pulled from any att (they share supplier via FK)
            vision_index.append((a.supplier.supplier_name, a.filename))

    return images, vision_index


def _extract_attachment_text(att: SupplierAttachment) -> str:
    """Dispatch on mime_type / extension. Returns '' for unsupported formats."""
    disk_path = _PROJECT_ROOT / att.stored_path
    if not disk_path.exists():
        return ""

    # Cache by (path, mtime) to avoid re-parsing unchanged files.
    try:
        mtime = disk_path.stat().st_mtime
    except OSError:
        return ""
    cache_key = (str(disk_path), mtime)
    if cache_key in _attachment_text_cache:
        return _attachment_text_cache[cache_key]

    name = (att.filename or "").lower()
    mime = (att.mime_type or "").lower()

    if mime == "application/pdf" or name.endswith(".pdf"):
        text = _extract_pdf_text(disk_path)
    elif "word" in mime or name.endswith((".docx",)):
        text = _extract_docx_text(disk_path)
    elif "spreadsheet" in mime or name.endswith((".xlsx", ".xlsm")):
        text = _extract_xlsx_text(disk_path)
    elif mime.startswith("text/") or name.endswith((".txt", ".csv", ".md", ".log")):
        text = _extract_plain_text(disk_path)
    else:
        text = ""   # images, archives, .doc (legacy) → list by name only

    text = text[:_ATTACHMENT_TEXT_BUDGET_PER_FILE]
    _attachment_text_cache[cache_key] = text
    return text


def _fmt_attachments(s: SavedSupplier, atts: list[SupplierAttachment]) -> str:
    """
    [ATTACHMENTS] block. Text-based files get their content inlined;
    images are marked "attached to AI vision" — they're sent separately
    as multimodal inline_data, not as text.
    """
    if not atts:
        return ""

    lines = ["[ATTACHMENTS — user-uploaded files for this supplier]"]
    total = 0
    images_shown = 0
    for a in atts:
        header = f"  📎 {a.filename}"
        if a.size_bytes:
            header += f" ({a.size_bytes // 1024} KB)"

        if _is_image_attachment(a):
            # Images are fed through Gemma vision — mark them so the AI
            # knows to look at the visual content it received alongside
            # this prompt. Respect the per-supplier cap so the hint here
            # matches what was actually sent.
            if images_shown < _MAX_IMAGES_PER_SUPPLIER:
                lines.append(header + " — attached to AI vision ✓")
                images_shown += 1
            else:
                lines.append(header + " — image not sent to AI (over limit)")
            continue

        text = _extract_attachment_text(a) if total < _ATTACHMENT_TEXT_BUDGET_TOTAL else ""
        lines.append(header)
        if text:
            remaining = _ATTACHMENT_TEXT_BUDGET_TOTAL - total
            snippet = text[:remaining]
            indented = "\n".join("      " + line for line in snippet.splitlines() if line.strip())
            lines.append(indented)
            total += len(snippet)
        else:
            lines.append("      (binary or unsupported format — AI can reference the filename but cannot read contents)")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Chat intent classifier (#1) — keyword-based with priority order.
# COMPARE must check first (otherwise "A vs B 哪个更便宜" gets routed to RANK
# because "便宜" trips the RANK keyword). Order = priority.
# ---------------------------------------------------------------------------

_INTENT_KEYWORDS: list[tuple[str, list[str]]] = [
    ("COMPARE", ["对比", "compare", " vs ", " versus ", "和.*哪个", "head-to-head"]),
    ("OPINION", ["为什么", "why ", "靠谱吗", "可信", "你觉得", "think", "your opinion",
                 "好不好", "怎么样",
                 # +Bug 1 fix: analysis / future / assessment intents
                 "分析", "未来", "前景", "评估", "评价", "看法",
                 "analyze", "future", "prospect", "outlook", "assessment", "review of"]),
    ("RANK",    ["最", "lowest", "highest", "cheapest", "safest", "best", "top "]),
    ("LOOKUP",  ["多少", "how much", "moq", "price", "lead time", "delivery",
                 "email", "phone", "联系", "交期", "报价", "价格"]),
    ("FIND",    ["哪家", "哪个", "which supplier", "which one", "who has", "who sells",
                 "find", "找一家", "找一个", "哪一家"]),
]


def _detect_query_language(query: str) -> str:
    """
    Decide whether to reply in Chinese or English. Mixed-language queries
    (a common case: "China Copper Sheet 未来前景") should be answered in
    Chinese because the user is Chinese-speaking and just used the
    supplier's English brand name verbatim.

    Detection: ANY of (≥2 Chinese chars) OR (>20% of chars are Chinese)
    triggers Chinese; otherwise English.
    """
    if not query:
        return "English"
    cn_chars = sum(1 for c in query if "一" <= c <= "鿿")
    total = len(query)
    if cn_chars >= 2:
        return "Chinese"
    if total > 0 and cn_chars / total > 0.20:
        return "Chinese"
    return "English"


def _classify_chat_intent(query: str) -> str:
    """
    Return one of: COMPARE / OPINION / RANK / LOOKUP / FIND.

    Falls back to FIND when no keyword hits — FIND's prompt is the safest
    default (1-line direct answer about which supplier matches), better
    than letting the LLM go report-mode again.
    """
    if not query:
        return "FIND"
    q = query.lower()
    for intent, kws in _INTENT_KEYWORDS:
        for kw in kws:
            # Treat regex-looking patterns separately so substring "和" doesn't
            # match every Chinese question with a comma.
            if ".*" in kw:
                if re.search(kw, q):
                    return intent
            else:
                if kw in q:
                    return intent
    return "FIND"


# Per-intent prompt overrides. We do NOT replace the existing 120-line
# rule body — we LAYER an intent-specific instruction block on top so the
# model has hard length / format constraints regardless of the FORMAT A/B/C
# JSON schema below.

_INTENT_PROMPT_OVERRIDE: dict[str, str] = {
    "FIND": (
        "INTENT: FIND — user wants to know which supplier(s) match a category "
        "or criterion.\n"
        "OUTPUT: maximum 3 short sentences. State the matching supplier(s) "
        "directly. If only one matches, say so plainly. If none match, say so "
        "and suggest expanding the search. Do NOT discuss non-matching "
        "suppliers. Do NOT write per-supplier paragraphs."
    ),
    "LOOKUP": (
        "INTENT: LOOKUP — user wants a specific fact (MOQ, lead time, email, "
        "price, etc.).\n"
        "OUTPUT: maximum 3 short sentences. If the fact exists in VERIFIED or "
        "ATTACHMENTS, quote it directly. If the fact is not in any data layer, "
        "say 'I don't have that information' — do not invent or estimate."
    ),
    "RANK": (
        "INTENT: RANK — user wants the best / lowest / highest / cheapest by "
        "some criterion.\n"
        "OUTPUT: maximum 3 short sentences. Name the top supplier with one "
        "short reason. If the criterion is ambiguous, default to value_score "
        "and state that assumption in one clause."
    ),
    "COMPARE": (
        "INTENT: COMPARE — user wants head-to-head comparison between "
        "specific suppliers.\n"
        "OUTPUT: 4-8 lines, bullet style (NOT a markdown table — those break "
        "in the renderer). Cover Country, Risk, Price, and one distinguishing "
        "field. End with a one-sentence verdict on which is preferable for the "
        "user's likely goal."
    ),
    "OPINION": (
        "INTENT: OPINION — user wants a judgment, assessment, or analysis.\n"
        "{VERDICT_BLOCK}\n"
        "Do NOT include placeholders, bracketed text, or curly-brace tokens "
        "in your output. Do NOT invent industry analysis from generic "
        "knowledge. Do NOT speculate beyond what's in the supplier data. "
        "Avoid hedging language ('might', 'could', 'perhaps')."
    ),
}


# Language-specific Deep Report redirect line. Injected into the OPINION
# prompt at request time via {DEEP_REPORT_LINE} substitution. Button name
# "AI Deep Report" stays untranslated so the user can find the actual button
# in the UI (the UI label is English).
_DEEP_REPORT_TEMPLATES: dict[str, str] = {
    "Chinese": "如需了解市场趋势和最新动态的深入分析，请点击 AI Deep Report 按钮。",
    "English": "For deeper analysis on market trends and recent news, click the AI Deep Report button.",
}


# ---------------------------------------------------------------------------
# OPINION verdict layer — deterministic decision logic for multi-supplier
# OPINION queries. Code-side computes the mode (RECOMMEND / HEDGE /
# INSUFFICIENT_DATA / SINGLE) plus the winner and reasons; the LLM is only
# responsible for writing the prose. Eligibility for RECOMMEND requires real
# evidence-based differentiation — the rules here are the guardrails that
# keep the LLM from inventing reasons on noisy data.
# ---------------------------------------------------------------------------

_RISK_RANK: dict[str, int] = {"low": 0, "medium": 1, "high": 2}

# Phone or email regex — same pattern used by modules.cleaner._CONTACT_RE.
_VERDICT_CONTACT_RE = re.compile(
    r"\+?\d[\d\s\-()]{8,15}"
    r"|[\w.+-]+@[\w-]+\.[\w.]+"
)

_VALUE_DIFF_THRESHOLD = 10.0   # min abs gap to count as a differentiator

# Keywords that signal the user is asking about data we structurally don't
# have (future / market / financial / news / industry analysis / growth).
# Mirrors the OPINION prompt's earlier description of "data not available".
_EXTERNAL_INSIGHT_KEYWORDS: list[str] = [
    "future", "outlook", "prospect", "prospects", "trend", "trends",
    "market", "financial", "finance", "news", "industry analysis", "growth",
    "未来", "前景", "趋势", "市场", "财务", "新闻", "行业分析", "增长",
]


def _has_contact(s) -> bool:
    """Phone or email present in the supplier's saved description."""
    return bool(_VERDICT_CONTACT_RE.search(s.description or ""))


def _requires_external_insight(query: str) -> bool:
    if not query:
        return False
    q = query.lower()
    return any(k in q for k in _EXTERNAL_INSIGHT_KEYWORDS)


def _risk_rank(s) -> int:
    return _RISK_RANK.get((s.risk_level or "").lower(), 99)


def _value_of(s) -> float:
    return float(s.value_score or 0.0)


def _pick_winner(suppliers: list):
    """
    Deterministic winner selection: risk (lower) > contact (has) > value
    (higher) > name (alphabetical, final tiebreak for stability).
    """
    # Stage 1: lowest risk
    min_r = min(_risk_rank(s) for s in suppliers)
    top = [s for s in suppliers if _risk_rank(s) == min_r]

    # Stage 2: prefer suppliers with contact info, but only narrow if at
    # least one in the current top set has it (otherwise everyone tied).
    if any(_has_contact(s) for s in top):
        top = [s for s in top if _has_contact(s)]

    # Stage 3: highest value
    max_v = max(_value_of(s) for s in top)
    top = [s for s in top if _value_of(s) == max_v]

    # Stage 4: alphabetical name — final stable tiebreak so result does
    # not flip on input order.
    return sorted(top, key=lambda s: (s.supplier_name or "").lower())[0]


def _compute_reasons(suppliers: list, winner) -> list[str]:
    """
    Return reason codes (subset of {'risk', 'contact', 'value'}, max 2)
    where the winner ACTUALLY occupies the favorable side AND a real diff
    exists across suppliers. Order is priority order: risk > contact > value.
    """
    reasons: list[str] = []

    risks = [_risk_rank(s) for s in suppliers]
    if _risk_rank(winner) == min(risks) and any(r != _risk_rank(winner) for r in risks):
        reasons.append("risk")

    contacts = [_has_contact(s) for s in suppliers]
    if _has_contact(winner) and any(not c for c in contacts):
        reasons.append("contact")

    values = [_value_of(s) for s in suppliers]
    if _value_of(winner) == max(values) and (max(values) - min(values)) >= _VALUE_DIFF_THRESHOLD:
        reasons.append("value")

    return reasons[:2]


_VERDICT_REASON_LABELS: dict[str, dict[str, str]] = {
    "Chinese": {
        "risk":    "风险更低",
        "contact": "提供可验证的联系方式",
        "value":   "价值评分明显更高",
    },
    "English": {
        "risk":    "lower risk",
        "contact": "verifiable contact information",
        "value":   "a notably higher value score",
    },
}


def _format_reasons_for_prompt(reasons: list[str], lang: str) -> str:
    labels = _VERDICT_REASON_LABELS.get(lang, _VERDICT_REASON_LABELS["English"])
    parts = [labels[r] for r in reasons if r in labels]
    if not parts:
        return ""
    if lang == "Chinese":
        return "且".join(parts)
    if len(parts) == 1:
        return parts[0]
    return " and ".join(parts)


def _render_verdict_line(winner_name: str, reasons: list[str], lang: str) -> str:
    """
    Pre-render the recommendation sentence in the user's language so the LLM
    can copy it verbatim. Avoids putting bilingual templates inside the
    prompt and getting the LLM to choose between them.
    """
    reasons_text = _format_reasons_for_prompt(reasons, lang)
    if lang == "Chinese":
        return f"建议优先选择 {winner_name}，因为其{reasons_text}。"
    return f"Recommend {winner_name} because it has {reasons_text}."


def _decide_verdict_mode(suppliers: list, query: str) -> dict:
    """
    Compute the verdict mode for an OPINION request.

    Returns dict with at least {"mode": ...}; for RECOMMEND mode also
    contains {"winner": <SavedSupplier>, "reasons": [str]}.

    Modes:
      SINGLE             — N == 1 (use existing single-supplier OPINION shape)
      INSUFFICIENT_DATA  — N >= 2 AND query touches external-insight topic
      HEDGE              — N >= 3, OR N == 2 with no real differentiator,
                           OR pick_winner produced no valid reasons
      RECOMMEND          — N == 2 with a real differentiator and ≥1 reason
    """
    n = len(suppliers)
    if n <= 1:
        return {"mode": "SINGLE"}

    if _requires_external_insight(query):
        return {"mode": "INSUFFICIENT_DATA"}

    if n >= 3:
        return {"mode": "HEDGE"}

    # N == 2: check for any real differentiator
    risks = [_risk_rank(s) for s in suppliers]
    contacts = [_has_contact(s) for s in suppliers]
    values = [_value_of(s) for s in suppliers]

    risk_diff    = len(set(risks)) > 1
    contact_diff = len(set(contacts)) > 1
    value_diff   = (max(values) - min(values)) >= _VALUE_DIFF_THRESHOLD

    if not (risk_diff or contact_diff or value_diff):
        return {"mode": "HEDGE"}

    winner = _pick_winner(suppliers)
    reasons = _compute_reasons(suppliers, winner)

    # Defensive: if pick_winner is correct this list is non-empty whenever
    # at least one diff exists, but degrade to HEDGE rather than emit a
    # reasonless recommendation if some future change breaks that property.
    if not reasons:
        return {"mode": "HEDGE"}

    return {"mode": "RECOMMEND", "winner": winner, "reasons": reasons}


# Per-mode body that fills the {VERDICT_BLOCK} slot in the OPINION prompt.
# SINGLE keeps the prior single-supplier behavior (≤3 sentences + Deep Report
# fallback). The multi-supplier modes IGNORE the 3-sentence cap and impose
# an exact line shape so the LLM cannot drift back to free-form paragraphs.
_VERDICT_BLOCKS: dict[str, str] = {
    "SINGLE": (
        "OUTPUT: maximum 3 short sentences.\n"
        "If the question requires data not available (e.g. future prospects, "
        "market trends, financial health, recent news, industry analysis, "
        "growth potential):\n"
        "- First, give a brief assessment based ONLY on available supplier "
        "data (1 sentence).\n"
        "- Then, append the following sentence EXACTLY (verbatim, do not "
        "translate the button name 'AI Deep Report'):\n"
        "    {DEEP_REPORT_LINE}\n"
        "The entire response MUST be in the required language."
    ),
    "RECOMMEND": (
        "OUTPUT: Multi-supplier mode. IGNORE the 3-sentence limit AND any "
        "later instruction asking for a paragraph per supplier.\n"
        "Produce EXACTLY N+1 short lines, separated by single newlines:\n"
        "  - Lines 1..N: ONE line per supplier, in the order they appear "
        "in the data, format '<supplier name>: <one short clause about "
        "risk level, contact information, or value score>'. Each line must "
        "be ONE clause — not multiple sentences, not a paragraph.\n"
        "  - Line N+1 (verdict): copy the following sentence VERBATIM, "
        "with no edits, no extra reasons, no rephrasing:\n"
        "        {VERDICT_LINE}\n"
        "Do NOT add extra sentences before, between, or after these N+1 "
        "lines. The entire response MUST be in the required language."
    ),
    "HEDGE": (
        "OUTPUT: Multi-supplier mode. IGNORE the 3-sentence limit AND any "
        "later instruction asking for a paragraph per supplier.\n"
        "Produce EXACTLY N+1 short lines, separated by single newlines:\n"
        "  - Lines 1..N: ONE line per supplier, in the order they appear "
        "in the data, format '<supplier name>: <one short clause about "
        "risk level, contact information, or value score>'. Each line must "
        "be ONE clause — not multiple sentences, not a paragraph.\n"
        "  - Line N+1 (verdict): a single short sentence stating the "
        "suppliers are comparable. If ONE concrete difference exists you "
        "may name it in one short clause; if no concrete difference "
        "exists, do NOT invent one — say the choice should depend on the "
        "user's project requirements. Do NOT recommend any specific "
        "supplier on this line.\n"
        "Do NOT add extra sentences before, between, or after these N+1 "
        "lines. The entire response MUST be in the required language."
    ),
    "INSUFFICIENT_DATA": (
        "OUTPUT: Output EXACTLY 2 lines, no more, no less. Do NOT list "
        "per-supplier details. Do NOT recommend any specific supplier.\n"
        "Line 1: a single sentence stating the question (future outlook / "
        "market trend / financial health / recent news / etc.) cannot be "
        "assessed from the current supplier data.\n"
        "Line 2 (verbatim, do not translate the button name 'AI Deep "
        "Report'):\n"
        "    {DEEP_REPORT_LINE}\n"
        "The entire response MUST be in the required language."
    ),
}


# ---------------------------------------------------------------------------
# Follow-up handling — when the current query is a context-dependent
# follow-up like "更具体一点" with no scope signal of its own, re-run the
# classifier and filters against the PREVIOUS user query so the inherited
# scope (intent + suppliers + category) carries forward. This avoids
# concatenating user strings (which would pollute classification) and
# avoids needing any session storage / frontend round-trip.
# ---------------------------------------------------------------------------

# Substrings whose presence in a SHORT query (<= _FOLLOWUP_MAX_LEN chars)
# indicates a follow-up to the previous turn. Substring (not exact match)
# so natural phrasings like "能给出我更多的细节么" / "你的判断是什么"
# are caught, not just the canonical "更具体一点" / "elaborate" forms.
_FOLLOWUP_SUBSTRINGS: set[str] = {
    # ask for more detail / specificity
    "更具体", "更详细", "更多细节", "更多的细节",
    "再具体", "再详细", "详细一点", "详细点",
    "more details", "be more specific", "elaborate", "tell me more",
    # ask for judgment / summary on the previous turn
    "你的判断", "你的想法", "你的看法",
    "总结一下", "概括一下", "总结",
    "字总结", "字的总结",          # "300字总结" / "100字的总结"
    "summarize", "in summary", "your verdict", "your take",
    # length adjustment ("再短一点" / "再长一点" / "shorter" / ...)
    "再短", "再长", "更短", "更长",
    "shorter", "longer",
    # generic continuation
    "继续", "再说说", "go on", "continue",
}

# Length cap: long queries usually carry their own scope (a new question)
# even when they happen to contain a follow-up substring like "你觉得".
# 20 chars cleanly covers "你的判断是什么" / "更多的细节么" / "总结一下"
# while excluding "你觉得中国供应商的价格怎么样" (which is a real new
# question that just contains a follow-up-shaped phrase).
_FOLLOWUP_MAX_LEN = 20


def _is_followup(query: str) -> bool:
    """
    True when the (short) query is a context-dependent follow-up like
    '能给出我更多的细节么' / 'your verdict?'. The caller still gates this
    behind a signal check (護栏 2: signal-first) so a query that
    introduces fresh scope is treated as a new question even if it
    contains a follow-up substring.
    """
    if not query:
        return False
    q = query.strip().lower()
    if len(q) > _FOLLOWUP_MAX_LEN:
        return False
    return any(s in q for s in _FOLLOWUP_SUBSTRINGS)


def _has_chat_signal(query: str) -> bool:
    """
    True if `query` carries any scope signal of its own — used as the
    "no signal" gate before deciding to inherit from the previous turn.

    Signals checked:
      - Non-default intent keyword (anything other than the FIND fallback)
      - Category keyword from _CHAT_CATEGORY_KEYWORDS
      - A name token >=4 chars not in _NAME_FILTER_GENERIC_TOKENS
        (cheap proxy for "the query names a specific supplier")
    """
    if not query:
        return False
    if _classify_chat_intent(query) != "FIND":
        return True
    q_lower = query.lower()
    for kws in _CHAT_CATEGORY_KEYWORDS.values():
        if any(k in q_lower for k in kws):
            return True
    # Distinctive name-shaped tokens. Restrict to tokens with at least one
    # ASCII letter — _NAME_FILTER_GENERIC_TOKENS is ASCII-only, and saved
    # supplier names today are ASCII. A pure-CJK token like "更多的细节"
    # would otherwise be mis-detected as "this query mentions a supplier
    # name", silently disabling follow-up inheritance for Chinese users.
    q_norm = _normalize_name(query)
    distinctive = {
        t for t in re.findall(r"\w+", q_norm)
        if len(t) >= 4
        and t not in _NAME_FILTER_GENERIC_TOKENS
        and re.search(r"[a-z]", t)
    }
    return bool(distinctive)


def _last_user_query(history) -> str | None:
    """
    Return the most recent USER message content from chat history.
    Robust to ordering and to any trailing assistant turns. Tolerant of
    both Pydantic ChatTurn objects (attr access) and plain dicts.
    """
    if not history:
        return None
    for msg in reversed(history):
        role = getattr(msg, "role", None)
        if role is None and isinstance(msg, dict):
            role = msg.get("role")
        if role != "user":
            continue
        content = getattr(msg, "content", None)
        if content is None and isinstance(msg, dict):
            content = msg.get("content")
        if content:
            return content
    return None


_FOLLOWUP_PROMPT_PREFIX_TEMPLATES: dict[str, str] = {
    "Chinese": (
        "FOLLOW-UP CONTEXT:\n"
        "用户在追问，希望对上一轮答案给出更详细的版本。不要重复同样的总结；"
        "请提供更具体的细节、解释或支持论据。\n\n"
    ),
    "English": (
        "FOLLOW-UP CONTEXT:\n"
        "The user is asking for a more detailed version of the previous "
        "answer. Do NOT repeat the same summary. Provide more specific "
        "details, explanations, or supporting points.\n\n"
    ),
}


# ---------------------------------------------------------------------------
# Explicit web-source request — when the user says "search online", "网上",
# "google it", etc., they're telling us they want web data, not a recap of
# saved fields. The existing _QUERY_DIMENSIONS layer triggers Serper based on
# WHICH dimension is being asked (capacity / certification / news / ...);
# this layer is orthogonal: it triggers Serper based on WHICH SOURCE the
# user explicitly requested. Keeping them separate so the two semantics
# don't collide as either list grows.
#
# Intentionally only complex multi-character phrases — single short words
# like "网上" (2 chars) substring-match into unrelated compounds (e.g.
# "网上海供应商" contains "网上" at position 0–1 even though the user means
# "上海"). Using "网上信息" / "网上的" / "网上搜" avoids that whole class
# of false positives.
# ---------------------------------------------------------------------------

_WEB_INTENT_KEYWORDS: list[str] = [
    # Chinese — compound phrases only (no bare 2-char "网上" / "网络")
    "网上信息", "网上的", "网上搜", "网上找", "网上有",
    "在网上", "来自网上", "从网上",
    "网络信息", "网络搜",
    "搜索一下", "google一下", "查一下网",
    # English — phrases / unambiguous tokens
    "online", "internet", "web search", "search the web",
    "look online", "google",
]


def _user_wants_web(query: str) -> bool:
    """
    True when the user explicitly asked for online / web information.
    Orthogonal to _match_dimensions (which checks WHICH dimension is asked
    about). This one checks WHICH SOURCE the user wants.
    """
    if not query:
        return False
    q = query.lower()
    return any(k in q for k in _WEB_INTENT_KEYWORDS)


# When _user_wants_web is True, the per-intent COMPARE / OPINION / RANK /
# LOOKUP / FIND template is REPLACED entirely by this one. We don't try to
# "override" with prepended warnings because empirically (see the
# verdict-layer rollout) the LLM follows the most specific template in the
# prompt, not the highest-priority one — so the only deterministic fix is
# to remove the conflicting template, not annotate it.
_WEB_SUMMARY_OVERRIDE_TEMPLATES: dict[str, str] = {
    "Chinese": (
        "INTENT: WEB SUMMARY — 用户明确请求基于网上信息的回答。\n"
        "OUTPUT 要求：\n"
        "- 用 4-8 句连续段落（不是 bullet list，不是字段表）。\n"
        "- 主体内容来自每家 supplier 的 [LIVE-WEB] 段落（[web-search] / "
        "[web-fetch] 数据）。\n"
        "- 每条网上事实必须带 [web-search] 或 [web-fetch] 源标签。\n"
        "- 不要把 Country / Risk / Price / value_score 这些本地字段作为主线"
        "——它们只作背景，必要时一笔带过。\n"
        "- 用户明示字数（如\"500 字\"）时优先满足；未明示给 4-8 句中长。\n"
        "- 若 [LIVE-WEB] 段落空缺或几乎没内容，明确告知用户暂未检索到足够"
        "的网上信息，不要编造。"
    ),
    "English": (
        "INTENT: WEB SUMMARY — user explicitly requested information from "
        "the web.\n"
        "OUTPUT requirements:\n"
        "- 4-8 sentences in flowing prose (NOT bullets, NOT a field list).\n"
        "- The body of the answer comes from each supplier's [LIVE-WEB] "
        "block ([web-search] / [web-fetch] data).\n"
        "- Cite every web fact with its [web-search] or [web-fetch] tag.\n"
        "- Do NOT lead with Country / Risk / Price / value_score — local "
        "fields are background only, mention them only if needed for "
        "context.\n"
        "- Match any user-specified length; otherwise aim for 4-8 sentences.\n"
        "- If the [LIVE-WEB] sections are empty or near-empty, say so "
        "plainly — do NOT invent web findings."
    ),
}


# Conditional tail — only attached when the user's question touches data
# we structurally don't have (price / MOQ / lead time), AND the intent is
# LOOKUP or RANK (so we're not redirecting an OPINION or COMPARE answer).

_TAIL_KEYWORDS = [
    "price", "cost", "quote", "moq",
    "lead time", "delivery",
    "交期", "价格", "报价", "成本",
    "how much", "多少钱",
    # Price-implying ranking words — "cheapest copper supplier" implies the
    # user wants the price answer, so route them to RFQ as well.
    "cheap", "cheaper", "cheapest", "expensive", "afford", "便宜",
]

_TAIL_TEXT = "For MOQ, lead time, and exact pricing, request a quote."


def _should_add_tail(query: str, intent: str) -> bool:
    if intent not in ("LOOKUP", "RANK"):
        return False
    q = (query or "").lower()
    return any(k in q for k in _TAIL_KEYWORDS)


# ---------------------------------------------------------------------------
# Chat helpers (中修): pre-filter by category intent + strip noisy source tags.
# ---------------------------------------------------------------------------
# Goal: when the user asks "which one is a copper supplier", we don't want
# the LLM to dutifully list all 4 saved suppliers and apologize for the
# 3 ACP ones. Backend filters first → LLM only narrates the matching set.
#
# We don't have a `category` column on SavedSupplier, so we match against
# supplier_name + description + url text. Order of dict matters: stainless
# before steel (stainless contains "steel"), ACP before aluminum (ACP
# contains "aluminum composite").

_CHAT_CATEGORY_KEYWORDS: dict[str, list[str]] = {
    "stainless": ["stainless", "不锈钢"],
    "acp":       ["acp", "aluminum composite", "aluminium composite", "铝塑板", "铝复合板"],
    "copper":    ["copper", "铜", "cuprum"],
    "aluminum":  ["aluminum", "aluminium", "铝"],
    "steel":     ["steel", "钢"],
    "brass":     ["brass", "黄铜"],
    "zinc":      ["zinc", "锌"],
    "titanium":  ["titanium", "钛"],
}


# ---------------------------------------------------------------------------
# Name-based filter (Bug 2 fix)
#
# When the user mentions a specific supplier by name in their question, we
# narrow to that supplier BEFORE category filter (which would over-include
# similarly-named peers). Three-tier match strategy:
#
#   Tier 1: normalized full name appears in normalized query
#   Tier 2: ≥2 distinctive tokens overlap (after dropping GENERIC_TOKENS)
#   Tier 3: exactly 1 distinctive token, but length ≥5 (e.g. "Aldura")
#
# Plus 3 safety nets:
#   1. Normalize both sides (strip punctuation + corp suffixes)
#   2. Score-based tie-break: only narrow when winner is uniquely top
#   3. Unicode-preserving normalize so Chinese supplier names don't get
#      stripped (most are English, but defensive).
# ---------------------------------------------------------------------------

# B2B / industry generic tokens — substrings shared by many suppliers, so
# a query containing only these isn't enough to identify ONE supplier.
_NAME_FILTER_GENERIC_TOKENS: set[str] = {
    "copper", "steel", "sheet", "sheets", "panel", "panels",
    "aluminum", "aluminium", "stainless", "metal", "metals",
    "factory", "manufacturer", "manufacturers", "manufacturing",
    "company", "co", "ltd", "llc", "inc", "corp", "corporation",
    "industries", "industry", "supplier", "suppliers", "supply",
    "china", "india", "vietnam", "korea", "japan", "german", "germany",
    "from", "and", "the", "of", "in", "for", "real", "best", "top",
}


def _normalize_name(text: str) -> str:
    """
    Lowercase + strip punctuation (keep alphanumeric + Chinese) +
    drop common corp suffixes + collapse whitespace.
    Used on BOTH supplier name and user query for fair comparison.
    """
    if not text:
        return ""
    t = text.lower()
    # Keep ASCII alphanumeric + space + CJK Unified Ideographs (Chinese chars)
    t = re.sub(r'[^a-z0-9一-鿿\s]', ' ', t)
    # Drop standalone corp suffixes
    t = re.sub(r'\b(co|ltd|llc|inc|corp|corporation)\b', ' ', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def _filter_by_supplier_name(query: str, suppliers: list) -> list | None:
    """
    If the query unambiguously names ONE supplier, return [that supplier].
    Otherwise return None (caller falls through to category filter).

    Returns None when 0 candidates OR multiple equally-strong candidates.
    Caller is responsible for skipping this whole step in COMPARE intent.
    """
    q_norm = _normalize_name(query)
    if not q_norm:
        return None

    q_tokens = set(re.findall(r'\w+', q_norm))

    candidates: list[tuple[object, int]] = []   # (supplier, score)

    for s in suppliers:
        name = (s.supplier_name or "").lower()
        if not name:
            continue
        name_norm = _normalize_name(name)
        if not name_norm:
            continue

        # --- Tier 1: full normalized name appears in query (highest score) ---
        if name_norm in q_norm:
            candidates.append((s, 100))
            continue

        # --- Tier 1b: name MINUS trailing suffix tokens in query --------
        # Handles "Acme Co, Suppliers, Ltd" matching when user types "Acme Co".
        # Only relevant for names of 4+ tokens so we don't reduce short names
        # to ambiguity.
        name_words = name_norm.split()
        matched_1b = False
        if len(name_words) >= 4:
            for drop in (1, 2, 3):
                if drop >= len(name_words):
                    break
                partial = " ".join(name_words[:-drop])
                if len(partial.split()) < 3:
                    break
                if partial in q_norm:
                    candidates.append((s, 80 - drop * 5))
                    matched_1b = True
                    break
        if matched_1b:
            continue

        # --- Tier 2 / 3: distinctive token overlap ----------------------
        name_tokens = set(re.findall(r'\w+', name_norm))
        overlap = (name_tokens & q_tokens) - _NAME_FILTER_GENERIC_TOKENS
        # Filter very short tokens (≥4 chars) to avoid spurious matches
        overlap = {t for t in overlap if len(t) >= 4}

        if len(overlap) >= 2:                          # Tier 2
            candidates.append((s, len(overlap) + 1))
        elif len(overlap) == 1 and any(len(t) >= 5 for t in overlap):
            # Tier 3: single distinctive token but length ≥5 — strong enough
            # for unique brand names like "Aldura"
            candidates.append((s, 2))

    if not candidates:
        return None

    # Tie-break: only narrow when winner is UNIQUELY top (≥1 ahead of #2).
    # If two candidates tie, return None — let category filter or LLM
    # disambiguate via the question.
    candidates.sort(key=lambda x: x[1], reverse=True)
    if len(candidates) == 1 or candidates[0][1] >= candidates[1][1] + 1:
        return [candidates[0][0]]
    return None


def _filter_by_category_intent(query: str, suppliers: list) -> list | None:
    """
    Pre-filter `suppliers` to those matching the product category implied
    by the user's question. Returns None when no category was detected
    (caller keeps the full list and adds no scope note).
    """
    q = (query or "").lower()
    matched_keywords: list[str] | None = None
    for _, kws in _CHAT_CATEGORY_KEYWORDS.items():
        if any(k in q for k in kws):
            matched_keywords = kws
            break

    if matched_keywords is None:
        return None

    def _hits(s) -> bool:
        haystack = " ".join([
            (s.supplier_name or "").lower(),
            (s.description   or "").lower(),
            (s.url           or "").lower(),
        ])
        return any(kw in haystack for kw in matched_keywords)

    return [s for s in suppliers if _hits(s)]


# Stripped from LLM output before showing to the user. We keep the raw
# answer in the response so future audit / tooltip features can recover
# the source attribution; this only cleans the visible text.
#
# Bug 3 fix: handles BOTH single tags `[rule-based]` AND comma/semicolon
# separated lists `[rule-based, AI-analysis]` that the LLM sometimes emits.
_SOURCE_TAG_NAMES = (
    r"(?:rule-based|AI-analysis|web-fetch|web-search|attachment|"
    r"user-verified|AI Deep Report)"
)
_SOURCE_TAG_RE = re.compile(
    r"\s*\[" + _SOURCE_TAG_NAMES
    + r"(?:\s*[,;]\s*" + _SOURCE_TAG_NAMES + r")*"
    + r"\]"
)


def _strip_source_tags(text: str) -> str:
    if not isinstance(text, str):
        return text
    return _SOURCE_TAG_RE.sub("", text)


def _strip_source_tags_in_structured(obj):
    """Walk the LLM JSON and clean every string field. Idempotent."""
    if isinstance(obj, str):
        return _strip_source_tags(obj)
    if isinstance(obj, list):
        return [_strip_source_tags_in_structured(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _strip_source_tags_in_structured(v) for k, v in obj.items()}
    return obj


@router.post("/ai-search")
def ai_search(req: AiSearchRequest, db: Session = Depends(get_db)):
    all_suppliers = db.query(SavedSupplier).filter(SavedSupplier.is_saved == True).all()  # noqa: E712
    if not all_suppliers:
        return {"answer": "You haven't saved any suppliers yet. Run an analysis and save some suppliers first.", "results": []}

    # --- Scope: Compare mode (user selected N) vs Explore mode ---
    if req.selected_ids:
        suppliers = [s for s in all_suppliers if s.id in req.selected_ids]
        compare_mode = True
        scope_note = (f"COMPARE MODE — user has SELECTED {len(suppliers)} specific "
                      f"suppliers. Focus ONLY on these and structure the answer "
                      f"head-to-head across all of them (even if some have 'no data').")
    else:
        suppliers = all_suppliers
        compare_mode = False
        scope_note = ("EXPLORE MODE — user has not selected specific suppliers "
                      "in the UI. Default scope is the full saved set, BUT if "
                      "the user names a single supplier in the question itself, "
                      "narrow focus to only that supplier (see SCOPE RULE).")

    if not suppliers:
        return {"answer": "No suppliers match the selection.", "results": []}

    query_lower = req.query.lower()

    # --- Follow-up inheritance ------------------------------------------
    # When the user says "能给出更多细节么" / "你的判断是什么" right after
    # an answer, the current query carries no scope signal of its own.
    # Re-run the classifier and filters against the PREVIOUS user query so
    # the inherited intent + filters carry forward, instead of silently
    # broadening back to "all saved suppliers" and intent=FIND.
    #
    # Order matters (護栏 2): signal check runs FIRST. A query that
    # introduces fresh scope (new category / new supplier name / explicit
    # intent keyword) is always treated as a new question, even when it
    # contains a follow-up substring like "总结".
    # NOTE: previously gated on `not compare_mode` so explicitly-selected
    # suppliers wouldn't have their scope overridden by a previous turn.
    # That gate is too aggressive: compare_mode locks the SUPPLIER set
    # (via selected_ids), but the follow-up mechanism is also what lets
    # web_mode / intent inheritance carry across turns. Allowing follow-up
    # in compare_mode does NOT change which suppliers are in scope (the
    # filters still skip when compare_mode is True), it only lets the
    # downstream logic see effective_query so e.g. _user_wants_web can
    # inherit from "根据网上的资料..." into "能继续总结的更详细一点么?".
    _has_sig = _has_chat_signal(req.query)
    _is_fu = (not _has_sig) and _is_followup(req.query)
    _prev_query = _last_user_query(req.history) if _is_fu else None
    followup_mode = bool(
        _is_fu
        and _prev_query
        and _prev_query.strip() != req.query.strip()
    )
    effective_query = _prev_query if followup_mode else req.query

    # --- Intent classifier (#1) — drives the per-intent output prompt ---
    # Order: COMPARE > OPINION > RANK > LOOKUP > FIND. Compare-mode (user
    # ticked specific suppliers) hard-overrides to COMPARE because that's
    # the explicit user signal.
    chat_intent = "COMPARE" if compare_mode else _classify_chat_intent(effective_query)
    _log.info(
        "[ai-search] intent=%s followup=%s query=%r effective=%r",
        chat_intent, followup_mode, req.query[:80], effective_query[:80],
    )

    # --- Language lock (Step 1: language detection) ----------------------
    # Detect language from the user's query and inject a HARD directive at
    # the TOP of scope_note (highest LLM attention position). This is
    # stronger than relying on phykawing's Rule #12 alone, which fails on
    # mixed-language queries like "China Copper Sheet ... 未来前景".
    detected_lang = _detect_query_language(req.query)
    scope_note = (
        f"REPLY STRICTLY IN {detected_lang.upper()}.\n"
        f"Do NOT switch language under any circumstance, even if the question "
        f"contains foreign brand names, product codes, or technical terms.\n"
        f"Every free-form string in your JSON response MUST be in {detected_lang}.\n\n"
        + scope_note
    )
    if followup_mode:
        scope_note = (
            _FOLLOWUP_PROMPT_PREFIX_TEMPLATES.get(
                detected_lang, _FOLLOWUP_PROMPT_PREFIX_TEMPLATES["English"]
            )
            + scope_note
        )
    _log.info("[ai-search] language=%s", detected_lang)

    # Debug counter exposed in response — temporary, helps verify filter wiring
    _debug_pre_filter_count = len(suppliers)

    # --- 中修: pre-filter by category intent (Bug fix for "list everything") ---
    # If the user's question clearly names a product category (e.g. "copper",
    # "ACP", "stainless"), narrow the supplier set BEFORE building the prompt.
    # Without this, the LLM dutifully discusses every saved supplier and
    # apologises for the non-matching ones — the "feels not smart" output.
    # Compare-mode is excluded because the user has already explicitly chosen
    # which suppliers to compare; respect that.
    _debug_filter_trace = {
        "compare_mode": compare_mode,
        "intent": chat_intent,
        "name_filter": "not_called",
        "category_filter": "not_called",
        "query_repr": repr(req.query),
    }
    skip_category_filter = False

    # ---- Bug 2 fix: name filter (highest priority, COMPARE excluded) ----
    # COMPARE mode means user is asking about multiple suppliers explicitly,
    # so narrowing to a single named one would be wrong. Same for the
    # explicit selected_ids (compare_mode flag).
    if not compare_mode and chat_intent != "COMPARE":
        _name_match = _filter_by_supplier_name(effective_query, suppliers)
        _debug_filter_trace["name_filter"] = (
            None if _name_match is None
            else f"{len(_name_match)} items: {[s.supplier_name for s in _name_match]}"
        )
        if _name_match:
            _log.info(
                "[ai-search] name filter narrowed %d -> 1 (%s)",
                len(suppliers), _name_match[0].supplier_name,
            )
            suppliers = _name_match
            scope_note += (
                "\n\nNOTE: The user is referring to a specific supplier by "
                "name. Focus ONLY on that supplier. Do NOT mention or "
                "compare any other suppliers."
            )
            skip_category_filter = True

    # ---- 中修: category filter (only if name filter didn't already narrow) ---
    if not compare_mode and not skip_category_filter:
        _filtered = _filter_by_category_intent(effective_query, suppliers)
        _debug_filter_trace["category_filter"] = (
            None if _filtered is None
            else f"{len(_filtered)} items: {[s.supplier_name for s in _filtered]}"
        )
        if _filtered is not None:
            if 0 < len(_filtered) < len(suppliers):
                _log.info(
                    "[ai-search] category pre-filter: %d -> %d suppliers",
                    len(suppliers), len(_filtered),
                )
                suppliers = _filtered
                scope_note += (
                    f"\n\nNOTE: Suppliers have been pre-filtered to "
                    f"{len(suppliers)} based on the product category in the "
                    f"user's question. ONLY discuss these suppliers — do "
                    f"NOT mention or explain excluded ones."
                )
            elif len(_filtered) == 0:
                scope_note += (
                    "\n\nNOTE: The user's question implied a product category "
                    "but none of the saved suppliers obviously matches it. "
                    "Answer over the full saved set and tell the user that "
                    "no exact match was found."
                )

    # --- Smart Layer-C trigger plan ---
    # Compare mode: all selected suppliers are candidates (user asked about them).
    # Explore mode: only top-3 by value_score to control API cost.
    if compare_mode:
        default_targets = suppliers
    else:
        default_targets = sorted(suppliers, key=lambda s: s.value_score or 0, reverse=True)[:3]

    enrich_targets: dict[int, SavedSupplier] = {}   # id → supplier (Tavily extract)
    search_targets: dict[int, SavedSupplier] = {}   # id → supplier (Serper search)

    for keywords, field in _match_dimensions(req.query):
        if field == "__enrich__":
            for s in default_targets:
                enrich_targets[s.id] = s
            continue
        if field is None:
            # VERIFIED doesn't cover this dimension. But if a fresh cached
            # deep_report already has the answer, we can skip the live web
            # search (strategy-B optimization) — deep_report is grounded
            # in web data and cheaper than a re-fetch.
            dr_section = _deep_report_section_for_keywords(keywords)
            for s in default_targets:
                if (dr_section
                        and _deep_report_is_fresh(s)
                        and _deep_report_has_content(s, dr_section)):
                    continue   # already covered
                search_targets[s.id] = s
            continue
        # VERIFIED field known → only search suppliers whose value is unreliable
        for s in default_targets:
            if not _primary_is_reliable(s, field, query_lower):
                search_targets[s.id] = s

    # --- Explicit web-source request override ---
    # When the user says "网上信息" / "search online" / "google it", force a
    # Serper search regardless of whether _QUERY_DIMENSIONS matched any
    # specific dimension keyword. This is the routing fix for queries like
    # "这些公司的网上信息" where the user explicitly wants web data but no
    # specific dimension keyword (capacity / certification / news / ...) is
    # present in the query.
    #
    # Compare-mode: search every selected supplier (user picked the scope).
    # Explore-mode: cap at top-3 by value_score, same as the dimension-driven
    # path, to control Serper costs.
    # Use effective_query (not req.query) so a follow-up like "300字总结"
    # or "能继续总结的更详细一点么?" inherits the web-mode flag from the
    # previous user turn the same way intent and filters do. Without this,
    # web_mode silently falls off across follow-up turns even though the
    # rest of conversation state was inherited correctly.
    _user_web_requested = _user_wants_web(effective_query) and bool(suppliers)
    if _user_web_requested:
        web_targets_for_user = (
            list(suppliers) if compare_mode else default_targets
        )
        for s in web_targets_for_user:
            search_targets[s.id] = s
        _log.info(
            "[ai-search] user explicitly requested web — adding %d supplier(s) to Serper queue",
            len(web_targets_for_user),
        )

    # --- Layer-C fetches (parallel, TTL-cached) ---
    # Runs all Tavily + Serper calls concurrently. Each call goes through
    # a 5-minute TTL cache keyed on url / (supplier_name + normalized query),
    # so rapid-fire repeat questions don't re-hit the paid APIs.
    t_layer_c = time.time()
    enriched_data, web_results = _fetch_layer_c_parallel(
        enrich_targets, search_targets, req.query
    )
    if enrich_targets or search_targets:
        _log.info(
            "[ai-search] Layer-C done in %.1fs  enrich=%d/%d  search=%d/%d",
            time.time() - t_layer_c,
            len(enriched_data), len(enrich_targets),
            len(web_results),  len(search_targets),
        )

    # --- Load all attachments for in-scope suppliers in one query ---
    supplier_ids = [s.id for s in suppliers]
    att_rows = (db.query(SupplierAttachment)
                  .filter(SupplierAttachment.supplier_id.in_(supplier_ids))
                  .order_by(SupplierAttachment.uploaded_at.desc())
                  .all()) if supplier_ids else []
    atts_by_supplier: dict[int, list[SupplierAttachment]] = {}
    for a in att_rows:
        atts_by_supplier.setdefault(a.supplier_id, []).append(a)

    # --- Collect image attachments for Gemma vision ---
    # Images are passed alongside the prompt via `call_model(... images=...)`.
    # _fmt_attachments will mark them in the text prompt so the AI knows
    # which image it's looking at when multiple are attached.
    ai_images, vision_index = _collect_images_for_ai(atts_by_supplier)
    if ai_images:
        _log.info("[ai-search] sending %d image(s) to Gemma vision", len(ai_images))

    # --- Build per-supplier four-layer blocks ---
    supplier_data = "\n\n".join(
        _format_supplier_block(s, enriched_data, web_results, atts_by_supplier.get(s.id))
        for s in suppliers
    )

    # Short, ordered index of the images the AI will see. Lets the model
    # tie each "image N" it sees to a specific supplier filename.
    vision_block = ""
    if vision_index:
        lines = ["\n\nIMAGES ATTACHED TO THIS PROMPT (in order):"]
        for i, (supplier_name, filename) in enumerate(vision_index, start=1):
            lines.append(f"  Image {i}: {filename}  →  supplier: {supplier_name}")
        vision_block = "\n".join(lines)

    # --- Conversation history ---
    history_block = ""
    if req.history:
        turns = []
        for t in req.history[-6:]:   # last 6 turns only
            role = "User" if t.role == "user" else "Assistant"
            turns.append(f"{role}: {t.content}")
        history_block = "\n\nCONVERSATION HISTORY (for context — the user may refer to earlier answers):\n" + "\n".join(turns)

    # When the user explicitly asked for web information, REPLACE the
    # per-intent template entirely with the WEB_SUMMARY template. We can't
    # just prepend an "override" warning because the per-intent template
    # contains specific shape rules (e.g. COMPARE: "Cover Country, Risk,
    # Price, and one distinguishing field") and empirically the LLM follows
    # the more specific instruction even when an earlier section says
    # "this is ABSOLUTE". Removing the conflicting template is the only
    # deterministic fix.
    verdict_decision: dict = {"mode": "SINGLE"}
    if _user_web_requested:
        intent_override = _WEB_SUMMARY_OVERRIDE_TEMPLATES.get(
            detected_lang, _WEB_SUMMARY_OVERRIDE_TEMPLATES["English"]
        )
        verdict_decision = {"mode": "WEB_SUMMARY"}
        _log.info("[ai-search] intent_override -> WEB_SUMMARY (user wants web)")
    else:
        intent_override = _INTENT_PROMPT_OVERRIDE.get(
            chat_intent, _INTENT_PROMPT_OVERRIDE["FIND"]
        )

        # --- OPINION verdict layer ---------------------------------
        # Code-side decides whether to recommend, hedge, or redirect
        # (single / multi-supplier branching is here, not in the
        # prompt). The LLM only writes prose; the verdict shape is
        # system-determined.
        if chat_intent == "OPINION":
            verdict_decision = _decide_verdict_mode(suppliers, effective_query)
            verdict_block = _VERDICT_BLOCKS.get(
                verdict_decision["mode"], _VERDICT_BLOCKS["SINGLE"]
            )
            if verdict_decision["mode"] == "RECOMMEND":
                verdict_block = verdict_block.replace(
                    "{VERDICT_LINE}",
                    _render_verdict_line(
                        verdict_decision["winner"].supplier_name or "",
                        verdict_decision["reasons"],
                        detected_lang,
                    ),
                )
            intent_override = intent_override.replace("{VERDICT_BLOCK}", verdict_block)
            _log.info(
                "[ai-search] verdict mode=%s n=%d reasons=%s",
                verdict_decision["mode"], len(suppliers),
                verdict_decision.get("reasons"),
            )

        # OPINION template carries a {DEEP_REPORT_LINE} token that
        # must be replaced with the language-matched redirect
        # sentence BEFORE the prompt reaches the model. Other intents
        # don't use the token; .replace() is a no-op for them.
        # SINGLE and INSUFFICIENT_DATA verdict blocks contain this
        # token; RECOMMEND and HEDGE do not.
        intent_override = intent_override.replace(
            "{DEEP_REPORT_LINE}",
            _DEEP_REPORT_TEMPLATES.get(detected_lang, _DEEP_REPORT_TEMPLATES["English"]),
        )

    prompt = f"""You are a supplier intelligence assistant.

{scope_note}

================================================================
{intent_override}

⚠ THIS INTENT BLOCK IS ABSOLUTE.
It overrides EVERY output-shape rule below, including:
  - the "≤3 sentences" length cap
  - the "give a paragraph per supplier" guidance in FORMAT B
  - any other verbosity / structure hint anywhere in this prompt
If the intent block above says "one short clause per supplier + one
verdict line", produce EXACTLY that — even if FORMAT B asks for
paragraphs. The FORMAT A/B/C JSON schemas are just envelopes; the
prose inside the chosen schema's text field MUST follow the intent
block's shape. Do NOT lengthen, shorten, restructure, or merge lines.
================================================================

Each supplier below is presented in up to six data layers:
  [VERIFIED]      user-entered — GROUND TRUTH.
  [SECONDARY-A]   rule-based scoring + anomaly detection.
  [SECONDARY-B]   AI-derived adjustments (optional, lower authority).
  [ATTACHMENTS]   files the user uploaded (quotes, spec sheets, photos).
                  Image attachments marked "attached to AI vision ✓" are
                  also sent to you as image inputs alongside this prompt;
                  read their contents directly when answering.
  [DEEP-RESEARCH] AI-synthesized research from a previous web investigation
                  (Serper + Tavily digest). ONLY fact-heavy sections are
                  included — subjective/opinion sections are intentionally
                  omitted so you do not re-cite a previous AI's opinions
                  as your own evidence.
  [LIVE-WEB]      fetched this turn from the supplier's website or Google.

{supplier_data}{vision_block}{history_block}

User question: {req.query}

================================================================
CRITICAL RULES — read before answering
================================================================

1. VERIFIED fields are GROUND TRUTH. The user entered them after real
   verification (received a quote, reviewed a sample, audited the factory).
   NEVER contradict a VERIFIED value with web data.

2. Source priority when facts disagree:
      VERIFIED  >  ATTACHMENTS  >  LIVE-WEB  >  DEEP-RESEARCH  >  SECONDARY-A  >  SECONDARY-B
   (Attachments are user-curated documents — trustworthy, second only to
    VERIFIED. LIVE-WEB is the freshest external data. DEEP-RESEARCH is
    cached AI-synthesis of earlier web data — useful for concrete facts
    like certification names and company info, but since it was itself
    produced by AI, treat it as second-hand evidence.)

3. NEVER hedge VERIFIED values. Do not say "approximately", "around",
   "reportedly", or "claims to be". Quote the exact user-provided value.

4. When VERIFIED and LIVE-WEB disagree, the VERIFIED value IS the answer.
   Mention the web value only as an informational note about the
   discrepancy (often just an outdated website).
   Example: "MOQ is 500 [user-verified]. Note: their public site lists
   1000 [web-search], which may be outdated."

5. Tag every fact with its source using these EXACT tags (preserve
   capitalisation and internal spaces):
      [user-verified]  [attachment]  [web-fetch]  [web-search]
      [AI Deep Report]  [rule-based]  [AI-analysis]

   When citing anything drawn from the [DEEP-RESEARCH] data block, the
   corresponding output tag is [AI Deep Report] (matching the button
   that originally produced the report).

6. If a dimension is missing across ALL layers, say "no data available" —
   do NOT invent, paraphrase, or guess.

   DEEP-RESEARCH anti-echo rule: only reference the [DEEP-RESEARCH]
   block for concrete facts (certification names, dates, numbers,
   company info). Cite them with the tag [AI Deep Report].
   Do NOT reuse the block's phrasing for subjective claims like
   "concerning complaints" or "cautious approach recommended" — those
   were AI opinions in a previous turn and must not be re-cited as
   evidence.

7. Do not move facts between suppliers. Each fact belongs only to the
   supplier under whose block it appeared.

8. SCOPE RULE (HIGHEST PRIORITY — overrides rule 9 below):

   Inspect the user's question for supplier names. If the question
   mentions ONE specific supplier by name (or by an unambiguous
   descriptor like "the Chinese one"), answer ONLY about that supplier.
   Do NOT volunteer information about the other suppliers.

   Examples of single-supplier questions:
     - "Alstrong's contact email"
     - "Best Aluminum Composite Panel Manufacturer in China 能找到图片么?"
     - "What are Reynobond's certifications?"
     - "the lowest risk one — what's their MOQ?"
   → Answer for that ONE supplier only.

   The mode-awareness rule below only applies when the question does
   NOT name a specific supplier.

9. Mode awareness (applies only when the question is NOT about a
   specific supplier per rule 8):
   - COMPARE MODE → structure the answer as head-to-head across ALL
     selected suppliers, even when a supplier has "no data" on a dimension.
   - EXPLORE MODE → pick the best match(es) from the full saved set and
     justify the shortlist.

10. Before answering a ranking question (lowest/highest/best X), scan all
    suppliers and handle ties: if multiple suppliers tie on the primary
    criterion (e.g. risk=0), break the tie by higher value_score then
    lower price. Never pick a tied supplier with a worse value_score.

11. If CONVERSATION HISTORY is provided, resolve references like
    "that supplier" or "the first one" from earlier turns.

12. LANGUAGE LOCK — reply strictly in the SAME LANGUAGE the user used in
    their question. If the question is in Chinese, every free-form string
    in your JSON (summary, reasons, tradeoffs, info, answer) MUST be in
    Chinese. Do not translate to English just because the prompt is in
    English. Source tags like [user-verified] stay literal in both languages.

================================================================
OUTPUT FORMAT — return ONLY a JSON object (no markdown fences)
================================================================

Classify the question first:
  "recommendation" — asking for ONE top pick / ranking / "which is best"
  "info"           — asking for specific facts, OR asking for an open-ended
                     opinion / assessment / overview ("what do you think",
                     "tell me about them", "你觉得怎么样", "介绍一下"). In
                     this case discuss EVERY in-scope supplier, not just one.
  "mixed"          — both recommend AND give info

FORMAT A — recommendation:
{{
  "type": "recommendation",
  "recommendation": {{
    "supplier_name": "name of top recommended supplier",
    "country": "country",
    "reasons": ["reason 1 [source-tag]", "reason 2 [source-tag]"],
    "tradeoffs": ["tradeoff 1"]
  }},
  "summary": "2-3 sentence answer citing specific data with [source-tag]s",
  "highlights": ["supplier_name_1", "supplier_name_2"]
}}

FORMAT B — info:
{{
  "type": "info",
  "answer": "PLAIN STRING with line breaks. Follow the EXACT shape and
             length mandated by the intent block at the top of this prompt
             (do not lengthen, restructure, or add per-supplier paragraphs
             beyond what the intent block specifies). Tag every fact with
             its [source-tag].",
  "highlights": ["supplier_name_1", "supplier_name_2"]
}}

FORMAT C — mixed:
{{
  "type": "mixed",
  "recommendation": {{
    "supplier_name": "…",
    "country": "…",
    "reasons": ["…"],
    "tradeoffs": ["…"]
  }},
  "info": "PLAIN STRING with line breaks and source tags.",
  "summary": "Short sentence tying it together.",
  "highlights": ["supplier_name_1"]
}}

Output ONLY the JSON object. No markdown, no extra text."""

    try:
        raw = call_model(prompt, images=ai_images if ai_images else None)
        if not raw:
            return {"answer": "AI is currently unavailable.", "structured": None, "results": [_to_dict(s) for s in suppliers]}

        import json
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        structured = json.loads(cleaned)

        # 中修: strip [rule-based]/[AI-analysis]/etc. from visible text.
        # Keep the original LLM output in `raw_answer` so auditing /
        # tooltip / debug features can recover source attribution later.
        raw_structured = structured
        clean_structured = _strip_source_tags_in_structured(structured)
        clean_summary = _strip_source_tags(clean_structured.get("summary", raw))

        # Conditional tail (#4) — append only when the question touched
        # data we structurally don't have AND the intent is LOOKUP/RANK
        # (not OPINION/COMPARE, where redirecting to RFQ is jarring).
        if _should_add_tail(req.query, chat_intent):
            tail = _TAIL_TEXT
            if clean_summary and not clean_summary.endswith(tail):
                clean_summary = clean_summary.rstrip(".? ") + ". " + tail
            # Also append to the structured 'answer' / 'info' / 'summary'
            # fields the frontend may render directly.
            for k in ("answer", "info", "summary"):
                v = clean_structured.get(k)
                if isinstance(v, str) and v and not v.endswith(tail):
                    clean_structured[k] = v.rstrip(".? ") + ". " + tail

        return {
            "answer":      clean_summary,
            "structured":  clean_structured,
            "raw_answer":  raw,
            "raw_structured": raw_structured,
            "intent":      chat_intent,
            "_debug_filter": {
                "pre":  _debug_pre_filter_count,
                "post": len(suppliers),
                "names": [s.supplier_name for s in suppliers],
                "trace": _debug_filter_trace,
            },
            "_debug_verdict": {
                "mode":    verdict_decision.get("mode"),
                "winner":  (verdict_decision["winner"].supplier_name
                            if verdict_decision.get("winner") is not None else None),
                "reasons": verdict_decision.get("reasons"),
            },
            "_debug_followup": {
                "active":          followup_mode,
                "has_signal":      _has_sig,
                "matched_followup": _is_fu,
                "prev_query":      _prev_query[:120] if _prev_query else None,
                "effective_query": effective_query[:120],
            },
            "_debug_web": {
                "user_requested":  _user_web_requested,
                "enrich_count":    len(enrich_targets),
                "search_count":    len(search_targets),
                "enrich_returned": len(enriched_data),
                "search_returned": len(web_results),
            },
            "results":     [_to_dict(s) for s in suppliers],
        }
    except json.JSONDecodeError:
        _log.warning("AI search JSON parse failed for query: %s", req.query[:100])
        return {"answer": "AI search response could not be parsed. Please try again.", "structured": None, "results": [_to_dict(s) for s in suppliers]}
    except Exception:
        _log.exception("AI search failed")
        return {"answer": "AI search failed. Please try again.", "structured": None, "results": [_to_dict(s) for s in suppliers]}


def _to_dict(s: SavedSupplier) -> dict:
    return {
        "id": s.id,
        "supplier_name": s.supplier_name,
        "country": s.country,
        "price_display": s.price_display,
        "price_usd": s.price_usd,
        "risk_level": s.risk_level,
        "risk_score": s.risk_score,
        "risk_reasons": s.risk_reasons,
        "value_score": s.value_score,
        "url": s.url,
        "description": s.description,
        "trust": s.trust,
        "anomalies": s.anomalies,
        "ai_adjustment": s.ai_adjustment,
        "notes": s.notes,
        "saved_at": s.saved_at.isoformat() if s.saved_at else None,
        # User primary decision data
        "decision_stage": s.decision_stage,
        "rating": s.rating,
        "tags": s.tags,
        "pros": s.pros,
        "cons": s.cons,
        "quoted_price": s.quoted_price,
        "quoted_currency": s.quoted_currency,
        "quoted_unit": s.quoted_unit,
        "moq": s.moq,
        "lead_time_days": s.lead_time_days,
        "payment_terms": s.payment_terms,
        "incoterms": s.incoterms,
        "sample_status": s.sample_status,
        "sample_quality": s.sample_quality,
        "factory_verified_via": s.factory_verified_via,
        "coating_confirmed": s.coating_confirmed,
        "core_material_confirmed": s.core_material_confirmed,
        "fire_rating_confirmed": s.fire_rating_confirmed,
        "reference_1": s.reference_1,
        "reference_2": s.reference_2,
        "reference_3": s.reference_3,
        "warranty_years": s.warranty_years,
        "next_action_date": s.next_action_date.isoformat() if s.next_action_date else None,
    }


# ---------------------------------------------------------------------------
# Attachments — user-uploaded files attached to a saved supplier
# ---------------------------------------------------------------------------

def _attachment_to_dict(a: SupplierAttachment) -> dict:
    return {
        "id":          a.id,
        "supplier_id": a.supplier_id,
        "filename":    a.filename,
        "mime_type":   a.mime_type,
        "size_bytes":  a.size_bytes,
        "uploaded_at": a.uploaded_at.isoformat() if a.uploaded_at else None,
    }


@router.get("/suppliers/{supplier_id}/attachments")
def list_attachments(supplier_id: int, db: Session = Depends(get_db)):
    """Return metadata for every attachment belonging to this supplier."""
    _get_active_supplier(db, supplier_id)

    rows = (db.query(SupplierAttachment)
              .filter(SupplierAttachment.supplier_id == supplier_id)
              .order_by(SupplierAttachment.uploaded_at.desc())
              .all())
    return [_attachment_to_dict(a) for a in rows]


@router.post("/suppliers/{supplier_id}/attachments")
async def upload_attachment(
    supplier_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """Accept a multipart file upload and store it under uploads/{supplier_id}/."""
    _get_active_supplier(db, supplier_id)

    safe_name = _sanitize_filename(file.filename or "file")

    # Extension allowlist — reject before writing any bytes to disk
    ext = Path(safe_name).suffix.lower()
    if ext not in ALLOWED_UPLOAD_EXTS:
        raise HTTPException(
            status_code=415,
            detail=f"File type '{ext or 'none'}' is not allowed. Allowed: {', '.join(sorted(ALLOWED_UPLOAD_EXTS))}",
        )

    supplier_dir = _UPLOAD_ROOT / str(supplier_id)
    supplier_dir.mkdir(parents=True, exist_ok=True)

    # Two-phase write: reserve a unique disk name, stream the file there,
    # then persist the DB row. A random prefix avoids collisions and makes
    # the directory resistant to guessing.
    unique_prefix = uuid.uuid4().hex[:10]
    disk_name = f"{unique_prefix}_{safe_name}"
    disk_path = supplier_dir / disk_name

    size = 0
    try:
        with disk_path.open("wb") as out:
            while True:
                chunk = await file.read(1024 * 1024)  # 1 MB chunks
                if not chunk:
                    break
                size += len(chunk)
                if size > MAX_UPLOAD_BYTES:
                    disk_path.unlink(missing_ok=True)
                    raise HTTPException(
                        status_code=413,
                        detail=f"File too large. Maximum size is {MAX_UPLOAD_BYTES // (1024 * 1024)} MB.",
                    )
                out.write(chunk)
    except HTTPException:
        raise
    except Exception:
        _log.exception("Upload stream error for supplier %d", supplier_id)
        disk_path.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail="Upload failed")

    # Derive MIME from our sanitized filename to avoid trusting client Content-Type
    mime, _ = mimetypes.guess_type(safe_name)

    att = SupplierAttachment(
        supplier_id=supplier_id,
        filename=safe_name,
        stored_path=str(disk_path.relative_to(_PROJECT_ROOT).as_posix()),
        mime_type=mime or "application/octet-stream",
        size_bytes=size,
    )
    db.add(att)
    db.commit()
    db.refresh(att)
    return _attachment_to_dict(att)


@router.get("/suppliers/{supplier_id}/attachments/{att_id}")
def download_attachment(supplier_id: int, att_id: int, db: Session = Depends(get_db)):
    """Serve the stored file as a forced download (Content-Disposition: attachment)."""
    _get_active_supplier(db, supplier_id)
    att = (db.query(SupplierAttachment)
             .filter(SupplierAttachment.id == att_id,
                     SupplierAttachment.supplier_id == supplier_id)
             .first())
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")

    disk_path = _PROJECT_ROOT / att.stored_path
    if not disk_path.exists():
        raise HTTPException(status_code=410, detail="File is missing from disk")

    mime, _ = mimetypes.guess_type(att.filename)
    safe_header_name = url_quote(att.filename)
    return FileResponse(
        disk_path,
        media_type=mime or "application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_header_name}"},
    )


@router.delete("/suppliers/{supplier_id}/attachments/{att_id}")
def delete_attachment(supplier_id: int, att_id: int, db: Session = Depends(get_db)):
    """Remove the DB row and best-effort unlink the file on disk."""
    _get_active_supplier(db, supplier_id)
    att = (db.query(SupplierAttachment)
             .filter(SupplierAttachment.id == att_id,
                     SupplierAttachment.supplier_id == supplier_id)
             .first())
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")

    disk_path = _PROJECT_ROOT / att.stored_path
    try:
        disk_path.unlink(missing_ok=True)
    except Exception as e:
        _log.warning("Could not unlink %s: %s", disk_path, e)

    db.delete(att)
    db.commit()
    return {"ok": True}
